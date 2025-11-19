from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.http import QueryDict
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET
from django.utils.dateparse import parse_date
from django.db.models import Count, F, ExpressionWrapper, DurationField, Avg, Prefetch
from django.conf import settings
from django.template.loader import render_to_string
import json
import io
import os
import base64
from datetime import datetime
import openpyxl
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import cm
try:
    from docx import Document
except ImportError:
    Document = None
try:
    from weasyprint import HTML
except Exception:
    HTML = None
import pandas as pd
import plotly.express as px
from collections import defaultdict
from django.utils.text import slugify
from django.utils.timezone import localtime, now
from django.urls import reverse

from rest_framework import viewsets
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.pagination import PageNumberPagination
from rest_framework.pagination import PageNumberPagination

from django.contrib.auth.models import User
from django.core.paginator import Paginator
from .models import (
    UserProfile,
    DiarioDigital,
    Categoria,
    LinkRelevante,
    LinkRedSocial,
    Articulo,
    Actividad,
    Roles,
    TipoActividad,
    Domicilio,
    Vehiculo,
    Empleador,
    Alias,
    Vinculo,
    HechoDestacado,
    Telefono,
    InformeIndividual,
    SolicitudInfo,
    HerramientaOSINT,
    HechoDelictivo,
    BandaCriminal,
    RedSocial,
    TvDigital,
    RadioDigital,
    EstadoLink,
    LinkTvDigital,
    LinkRadioDigital,
    InformeBandaCriminal,
    JerarquiaPrincipal,
)
from .serializers import (
    UserProfileSerializer,
    DiarioDigitalSerializer,
    CategoriaSerializer,
    LinkRelevanteSerializer,
    LinkRedSocialSerializer,
    ArticuloSerializer,
    ActividadSerializer,
    InformeIndividualSerializer,
    DomicilioSerializer,
    VehiculoSerializer,
    EmpleadorSerializer,
    AliasSerializer,
    VinculoSerializer,
    HechoDestacadoSerializer,
    UserSerializer,
    ArticuloDetailSerializer,
    RedSocialSerializer,
    TvDigitalSerializer,
    RadioDigitalSerializer,
    LinkTvDigitalSerializer,
    LinkRadioDigitalSerializer,
)
from .forms import (
    InformeIndividualForm,
    CustomAuthenticationForm,
    CategoriaForm,
    DiarioForm,
    ArticuloForm,
    SolicitudInfoForm,
    SolicitudInfoRespuestaForm,
    HerramientaOSINTForm,
    RedSocialForm,
    TvDigitalForm,
    RadioDigitalForm,
    HechoDelictivoForm,
    BandaCriminalForm,
    InformeBandaCriminalForm,
)

# Asegúrate de que esta utilidad exista en tu proyecto
from app.utils.actividad import log_actividad


from urllib.parse import urlparse

from django.contrib import messages


import logging

logger = logging.getLogger(__name__)

WEASYPRINT_ERROR_MESSAGE = (
    "No se puede generar el PDF porque las dependencias de WeasyPrint no están instaladas en el servidor."
)


def _weasyprint_unavailable_response(request, fallback_url=None):
    messages.error(request, WEASYPRINT_ERROR_MESSAGE)
    if fallback_url:
        return redirect(fallback_url)
    return HttpResponse(WEASYPRINT_ERROR_MESSAGE, status=503)

# -------------------------- AUTENTICACIÓN --------------------------


def login_view(request):
    if request.method == "POST":
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            log_actividad(request, TipoActividad.LOGIN, "Inicio de sesión exitoso")
            rol = getattr(user.userprofile, "rol", None)
            if rol == Roles.PRENSA:
                return redirect("prensa")
            elif rol == Roles.CLASIFICACION:
                return redirect("clasificacion")
            elif rol == Roles.REDACCION:
                return redirect("redaccion")
            elif rol == Roles.ADMIN:
                return redirect("actividad")
            elif rol == Roles.INFORMES:
                return redirect("informes")
            elif rol == Roles.GERENCIA:
                return redirect("estadisticas")
            elif rol == Roles.GERENTE_PRODUCCION:
                return redirect("configuraciones")
            elif rol == Roles.CLIENTE:
                return redirect("consulta_informes")
            else:
                return render(request, "403.html", status=403)
        else:
            # Si el login falla, loguear el intento fallido
            log_actividad(request, TipoActividad.LOGIN, "Intento de inicio de sesión fallido")
    else:
        form = CustomAuthenticationForm()
    return render(request, "login.html", {"form": form})


def logout_view(request):
    log_actividad(request, TipoActividad.LOGOUT, "Cierre de sesión exitoso")
    logout(request)
    return redirect("login")


# -------------------------- PANELES WEB --------------------------


@login_required
def prensa_view(request):
    if request.user.userprofile.rol in [Roles.PRENSA, Roles.ADMIN]:
        solicitudes_pendientes = SolicitudInfo.objects.filter(
            fecha_respuesta__isnull=True
        ).count()
        
        # AGREGADO: Cargar herramientas OSINT para la pestaña correspondiente
        herramientas_osint = HerramientaOSINT.objects.all()
        
        return render(
            request,
            "prensa.html",
            {
                "solicitudes_pendientes": solicitudes_pendientes,
                "herramientas_osint": herramientas_osint,  # AGREGADO: Pasar al contexto
            },
        )
    return render(request, "403.html", status=403)


@login_required
def redaccion_view(request):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.PRENSA, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitudes_usuario = (
        SolicitudInfo.objects.select_related("respondido_por")
        .prefetch_related("articulos")
        .filter(usuario_creador=request.user)
        .order_by("-fecha_creacion")
    )
    return render(
        request,
        "redaccion.html",
        {
            "solicitudes_usuario": solicitudes_usuario,
        },
    )


def _bandas_info_payload():
    def persona_payload(persona):
        alias = list(persona.alias.values_list("nombre", flat=True))
        telefonos = list(persona.telefono.values_list("numero", flat=True))
        return {
            "id": persona.id,
            "nombre": f"{persona.apellido}, {persona.nombre}".strip(", "),
            "documento": persona.documento or "",
            "rol": persona.get_rol_display() if persona.rol else "",
            "situacion": persona.get_situacion_display() if persona.situacion else "",
            "actividad": persona.actividad or "",
            "banda": persona.banda or "",
            "alias": alias,
            "telefonos": telefonos,
        }

    personas_prefetch = InformeIndividual.objects.prefetch_related("alias", "telefono")
    bandas = (
        BandaCriminal.objects.all()
        .prefetch_related(
            "bandas_aliadas",
            "bandas_rivales",
            Prefetch("lideres", queryset=personas_prefetch, to_attr="prefetched_lideres"),
            Prefetch("miembros", queryset=personas_prefetch, to_attr="prefetched_miembros"),
        )
        .order_by("nombres")
    )

    datos = []
    for banda in bandas:
        lideres = getattr(banda, "prefetched_lideres", banda.lideres.all())
        miembros = getattr(banda, "prefetched_miembros", banda.miembros.all())
        lugartenientes = [
            miembro
            for miembro in miembros
            if (miembro.rol or "").lower() == "lugarteniente"
        ]
        if not lugartenientes:
            lideres_ids = {miembro.pk for miembro in lideres}
            lugartenientes = [
                miembro for miembro in miembros if miembro.pk not in lideres_ids
            ]
        datos.append(
            {
                "id": banda.id,
                "nombre": banda.nombre_principal,
                "zonas": banda.zonas_influencia_detalle,
                "lideres": [persona_payload(persona) for persona in lideres],
                "lugartenientes": [persona_payload(persona) for persona in lugartenientes],
                "miembros": [persona_payload(persona) for persona in miembros],
                "bandas_aliadas": [
                    {"id": aliada.id, "nombre": aliada.nombre_principal}
                    for aliada in banda.bandas_aliadas.all()
                ],
                "bandas_rivales": [
                    {"id": rival.id, "nombre": rival.nombre_principal}
                    for rival in banda.bandas_rivales.all()
                ],
            }
        )
    return datos


def _informes_banda_queryset():
    return (
        InformeBandaCriminal.objects.select_related("banda")
        .prefetch_related(
            "bandas_aliadas",
            "bandas_rivales",
            "jerarquias__miembro",
        )
        .order_by("-creado_en")
    )


@login_required
def informe_banda_crear_view(request):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = InformeBandaCriminalForm(request.POST)
        if form.is_valid():
            informe = form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Informe de banda criminal creado (ID {informe.pk}).",
            )
            messages.success(request, "Informe de banda creado correctamente.")
            return redirect("redaccion")
        messages.error(
            request,
            "Revisá los campos del formulario. Hay errores que deben corregirse.",
        )
    else:
        form = InformeBandaCriminalForm()

    return render(
        request,
        "form_informe_banda.html",
        {
            "form": form,
            "bandas_info": json.dumps(_bandas_info_payload(), ensure_ascii=False),
            "informes": _informes_banda_queryset(),
            "titulo_pagina": "Nuevo informe de banda criminal",
            "descripcion_pagina": "Completá la información solicitada para registrar el informe.",
            "volver_url": reverse("redaccion"),
        },
    )


@login_required
def informe_banda_detalle_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    jerarquias_prefetch = Prefetch(
        "jerarquias",
        queryset=JerarquiaPrincipal.objects.select_related("miembro").prefetch_related(
            "miembro__alias",
            "miembro__telefono",
        ),
    )
    informe = get_object_or_404(
        InformeBandaCriminal.objects.select_related("banda").prefetch_related(
            "bandas_aliadas",
            "bandas_rivales",
            jerarquias_prefetch,
        ),
        pk=pk,
    )
    lideres = [
        jerarquia.miembro
        for jerarquia in informe.jerarquias.all()
        if jerarquia.rol == JerarquiaPrincipal.Rol.LIDER
    ]
    lugartenientes = [
        jerarquia.miembro
        for jerarquia in informe.jerarquias.all()
        if jerarquia.rol == JerarquiaPrincipal.Rol.LUGARTENIENTE
    ]
    return render(
        request,
        "informe_banda_detalle.html",
        {
            "informe": informe,
            "lideres": lideres,
            "lugartenientes": lugartenientes,
        },
    )


@login_required
def informe_banda_editar_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    informe = get_object_or_404(InformeBandaCriminal, pk=pk)
    if request.method == "POST":
        form = InformeBandaCriminalForm(request.POST, instance=informe)
        if form.is_valid():
            form.save()
            messages.success(request, "Informe de banda actualizado correctamente.")
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Informe de banda criminal actualizado (ID {informe.pk}).",
            )
            return redirect("detalle_informe_banda", pk=informe.pk)
        messages.error(
            request,
            "Revisá los campos del formulario. Hay errores que deben corregirse.",
        )
    else:
        form = InformeBandaCriminalForm(instance=informe)

    return render(
        request,
        "form_informe_banda.html",
        {
            "form": form,
            "bandas_info": json.dumps(_bandas_info_payload(), ensure_ascii=False),
            "titulo_pagina": "Editar informe de banda criminal",
            "descripcion_pagina": f"Banda: {informe.banda.nombre_principal}",
            "volver_url": reverse("detalle_informe_banda", args=[informe.pk]),
        },
    )


@login_required
def informe_banda_eliminar_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    informe = get_object_or_404(InformeBandaCriminal, pk=pk)
    if request.method == "POST":
        nombre_banda = informe.banda.nombre_principal
        informe.delete()
        messages.success(
            request,
            f"Informe de la banda {nombre_banda or 'sin nombre'} fue eliminado.",
        )
        log_actividad(
            request,
            TipoActividad.OTRO,
            f"Informe de banda criminal eliminado (ID {pk}).",
        )
        return redirect("crear_informe_banda")
    messages.error(request, "Acción inválida para eliminar el informe.")
    return redirect("detalle_informe_banda", pk=pk)


@login_required
def informe_banda_exportar_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)
    if Document is None:
        messages.error(
            request,
            "No es posible generar el documento porque falta la dependencia python-docx. Instalála y volvé a intentarlo.",
        )
        return redirect("detalle_informe_banda", pk=pk)

    jerarquias_prefetch = Prefetch(
        "jerarquias",
        queryset=JerarquiaPrincipal.objects.select_related("miembro").prefetch_related(
            "miembro__alias",
            "miembro__telefono",
        ),
    )
    informe = get_object_or_404(
        InformeBandaCriminal.objects.select_related("banda").prefetch_related(
            "bandas_aliadas",
            "bandas_rivales",
            jerarquias_prefetch,
        ),
        pk=pk,
    )

    banda = informe.banda
    jerarquias = list(informe.jerarquias.all())
    bandas_aliadas = list(banda.bandas_aliadas.all())
    bandas_rivales = list(banda.bandas_rivales.all())
    miembros_totales = list(banda.miembros.all())

    lideres = [
        jerarquia.miembro
        for jerarquia in jerarquias
        if jerarquia.rol == JerarquiaPrincipal.Rol.LIDER and jerarquia.miembro
    ]
    lugartenientes = [
        jerarquia.miembro
        for jerarquia in jerarquias
        if jerarquia.rol == JerarquiaPrincipal.Rol.LUGARTENIENTE and jerarquia.miembro
    ]

    doc = Document()
    doc.add_heading("Informe de Banda", 0)
    doc.add_paragraph(
        f"Banda: {banda.nombre_principal or 'Sin nombre'} · Generado el {localtime(now()).strftime('%d/%m/%Y %H:%M')}"
    )

    def add_anexo(numero, titulo, contenido, bullet=False):
        encabezado = f"Anexo {numero} - {titulo}"
        doc.add_heading(encabezado, level=1)
        if not contenido:
            doc.add_paragraph("Sin información disponible.")
            return
        if isinstance(contenido, (list, tuple)):
            for linea in contenido:
                if not linea:
                    continue
                if bullet:
                    doc.add_paragraph(linea, style="List Bullet")
                else:
                    doc.add_paragraph(linea)
        else:
            doc.add_paragraph(contenido)

    zonas = banda.zonas_influencia_detalle
    zonas_listado = [
        zona
        for zona in [
            ", ".join(
                list(
                    filter(
                        None,
                        [
                            z.get("barrio") or "",
                            z.get("localidad") or "",
                            z.get("ciudad") or "",
                            z.get("provincia") or "",
                        ],
                    )
                )
            )
            for z in zonas or []
        ]
        if zona
    ]

    miembros_lineas = []
    if lideres:
        miembros_lineas.append("Líderes identificados:")
        for lider in lideres:
            miembros_lineas.append(
                f"  - {lider} · Documento: {lider.documento or '-'} · Situación: {lider.get_situacion_display() or '-'}"
            )
    if lugartenientes:
        miembros_lineas.append("Lugartenientes identificados:")
        for lugarteniente in lugartenientes:
            miembros_lineas.append(
                f"  - {lugarteniente} · Documento: {lugarteniente.documento or '-'} · Situación: {lugarteniente.get_situacion_display() or '-'}"
            )
    miembros_sin_jerarquia = [
        miembro
        for miembro in miembros_totales
        if miembro.pk not in {m.pk for m in lideres + lugartenientes if m}
    ]
    if miembros_sin_jerarquia:
        miembros_lineas.append(
            "Miembros registrados sin jerarquía: "
            + ", ".join(str(miembro) for miembro in miembros_sin_jerarquia)
        )
    if not miembros_lineas:
        miembros_lineas = ["No se registraron miembros ni jerarquías para esta banda."]

    zonas_contenido = zonas_listado or ["No se cargaron zonas de influencia para la banda."]
    aliadas_contenido = (
        [f"- {aliada.nombre_principal or '-'}" for aliada in bandas_aliadas]
        if bandas_aliadas
        else ["Sin alianzas registradas."]
    )
    rivales_contenido = (
        [f"- {rival.nombre_principal or '-'}" for rival in bandas_rivales]
        if bandas_rivales
        else ["Sin rivales registrados."]
    )
    conclusion_texto = (
        informe.conclusion_relevante or "No se registró una conclusión relevante."
    )
    evolucion_texto = (
        informe.posible_evolucion or "No se registró una posible evolución."
    )

    anexo1_contenido = []
    anexo1_contenido.append("Miembros y jerarquías:")
    anexo1_contenido.extend(miembros_lineas)
    anexo1_contenido.append("")
    anexo1_contenido.append("Zonas de influencia:")
    anexo1_contenido.extend(zonas_contenido)
    anexo1_contenido.append("")
    anexo1_contenido.append("Bandas aliadas:")
    anexo1_contenido.extend(aliadas_contenido)
    anexo1_contenido.append("")
    anexo1_contenido.append("Bandas rivales:")
    anexo1_contenido.extend(rivales_contenido)
    anexo1_contenido.append("")
    anexo1_contenido.append("Conclusión relevante:")
    anexo1_contenido.append(conclusion_texto)
    anexo1_contenido.append("")
    anexo1_contenido.append("Probable evolución:")
    anexo1_contenido.append(evolucion_texto)

    add_anexo(1, "Resumen ejecutivo", anexo1_contenido)

    introduccion_contenido = (
        informe.introduccion_descripcion
        or "No se registró una descripción de introducción para este informe."
    )
    add_anexo(2, "Introducción", introduccion_contenido)

    antecedentes_contenido = []
    antecedentes_contenido.append("Hechos delictivos asociados a la banda:")
    hechos_relacionados = (
        banda.hechos_delictivos.select_related("articulo")
        .prefetch_related("autor")
        .order_by("-fecha")
    )
    if hechos_relacionados:
        for hecho in hechos_relacionados:
            autores_texto = ", ".join([str(autor) for autor in hecho.autor.all()])
            antecedentes_contenido.append(
                f"- {hecho.fecha.strftime('%d/%m/%Y')} · {hecho.get_categoria_display() or 'Sin categoría'} · {hecho.ubicacion_texto} · Calificación {hecho.get_calificacion_display()} · Autores: {autores_texto or ('Autor no identificado' if hecho.autor_desconocido else 'Sin autores')}"
            )
    else:
        antecedentes_contenido.append("No se registraron hechos delictivos asociados.")
    if informe.antecedentes:
        antecedentes_contenido.append("")
        antecedentes_contenido.append("Antecedentes personalizados:")
        for antecedente in informe.antecedentes:
            titulo = antecedente.get("titulo") if isinstance(antecedente, dict) else ""
            descripcion = (
                antecedente.get("descripcion") if isinstance(antecedente, dict) else ""
            )
            linea = f"{titulo or 'Sin título'}: {descripcion or 'Sin descripción'}"
            antecedentes_contenido.append(linea.strip())
    add_anexo(3, "Antecedentes", antecedentes_contenido)

    desarrollo_contenido = informe.desarrollo_contenido or "No se registró contenido para el desarrollo."
    desarrollo_titulo = informe.desarrollo_titulo or "Desarrollo"
    add_anexo(4, desarrollo_titulo, desarrollo_contenido)

    hechos_lineas = []
    if hechos_relacionados:
        for hecho in hechos_relacionados:
            titulo = f"{hecho.get_categoria_display() or 'Sin categoría'} - {hecho.fecha.strftime('%d/%m/%Y')}"
            hechos_lineas.append(titulo)
            hechos_lineas.append(f"Fecha del hecho: {hecho.fecha.strftime('%d/%m/%Y')}")
            hechos_lineas.append(f"Lugar: {hecho.ubicacion_texto}")
            descripcion = hecho.descripcion.strip() if hecho.descripcion else ""
            hechos_lineas.append(f"Descripción: {descripcion or 'Sin descripción registrada.'}")
            hechos_lineas.append("")
    else:
        hechos_lineas.append("No se registraron hechos delictivos asociados a la banda.")
    add_anexo(5, "Hechos relevantes", hechos_lineas)

    add_anexo(
        6,
        "Conclusiones",
        informe.conclusion_relevante
        or "No se registraron conclusiones en este informe.",
    )

    fichas = []
    personas = []
    personas.extend(lideres)
    personas.extend(lugartenientes)
    for persona in personas:
        alias = ", ".join(persona.alias.values_list("nombre", flat=True))
        telefonos = ", ".join(
            [str(numero) for numero in persona.telefono.values_list("numero", flat=True)]
        )
        fichas.append(
            f"{persona} · Documento: {persona.documento or '-'} · Rol: {persona.get_rol_display() or '-'} · Alias: {alias or 'Sin alias'} · Teléfonos: {telefonos or 'Sin teléfonos'}"
        )
    if not fichas:
        fichas.append("No hay fichas individuales asociadas al informe.")
    add_anexo(7, "Anexo fichas individuales", fichas, bullet=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_slug = slugify(banda.nombre_principal or "banda")
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response[
        "Content-Disposition"
    ] = f'attachment; filename="Informe_Banda_{nombre_slug}_{informe.pk}.docx"'
    return response


@login_required
def hechos_delictivos_view(request):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = HechoDelictivoForm(request.POST, user=request.user)
        if form.is_valid():
            hecho = form.save(commit=False)
            hecho.creado_por = request.user
            hecho.save()
            form.save_m2m()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Hecho delictivo creado desde panel (ID {hecho.pk}).",
            )
            messages.success(request, "Hecho delictivo registrado correctamente.")
            return redirect("hechos_delictivos")
        messages.error(
            request,
            "No se pudo registrar el hecho delictivo. Verificá los datos ingresados.",
        )
    else:
        form = HechoDelictivoForm(user=request.user)

    hechos = (
        HechoDelictivo.objects.filter(creado_por=request.user)
        .select_related("creado_por", "articulo")
        .prefetch_related("autor", "noticias")
        .order_by("-fecha")
    )

    return render(
        request,
        "hechos_delictivos.html",
        {
            "form": form,
            "hechos": hechos,
        },
    )


@login_required
def bandas_criminales_view(request):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    filtros = {
        "nombre": request.GET.get("nombre", "").strip(),
        "zona": request.GET.get("zona", "").strip(),
        "aliada": request.GET.get("aliada", "").strip(),
        "rival": request.GET.get("rival", "").strip(),
    }

    bandas_queryset = BandaCriminal.objects.all().prefetch_related(
        "lideres", "miembros", "bandas_aliadas", "bandas_rivales"
    )

    if filtros["nombre"]:
        bandas_queryset = bandas_queryset.filter(nombres__icontains=filtros["nombre"])
    if filtros["zona"]:
        bandas_queryset = bandas_queryset.filter(
            zonas_influencia__icontains=filtros["zona"]
        )
    if filtros["aliada"]:
        bandas_queryset = bandas_queryset.filter(
            bandas_aliadas__nombres__icontains=filtros["aliada"]
        )
    if filtros["rival"]:
        bandas_queryset = bandas_queryset.filter(
            bandas_rivales__nombres__icontains=filtros["rival"]
        )

    bandas_queryset = bandas_queryset.distinct()
    bandas = sorted(
        bandas_queryset,
        key=lambda banda: (banda.nombre_principal or "").lower(),
    )

    if request.method == "POST":
        form = BandaCriminalForm(request.POST)
        if form.is_valid():
            banda = form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Banda criminal creada (ID {banda.pk}).",
            )
            messages.success(request, "Banda criminal creada correctamente.")
            return redirect("bandas_criminales")
        messages.error(
            request,
            "No se pudo crear la banda criminal. Revisá los datos ingresados.",
        )
    else:
        form = BandaCriminalForm()

    return render(
        request,
        "bandas_criminales.html",
        {
            "form": form,
            "bandas": bandas,
            "modo_edicion": False,
            "form_action": reverse("bandas_criminales"),
            "filtros": filtros,
        },
    )


@login_required
def banda_criminal_editar_view(request, pk):
    banda = get_object_or_404(BandaCriminal, pk=pk)

    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    filtros = {
        "nombre": request.GET.get("nombre", "").strip(),
        "zona": request.GET.get("zona", "").strip(),
        "aliada": request.GET.get("aliada", "").strip(),
        "rival": request.GET.get("rival", "").strip(),
    }

    if request.method == "POST":
        form = BandaCriminalForm(request.POST, instance=banda)
        if form.is_valid():
            banda = form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Banda criminal actualizada (ID {banda.pk}).",
            )
            messages.success(request, "Banda criminal actualizada correctamente.")
            return redirect("bandas_criminales")
        messages.error(
            request,
            "No se pudo actualizar la banda criminal. Revisá los datos ingresados.",
        )
    else:
        form = BandaCriminalForm(instance=banda)

    bandas_queryset = BandaCriminal.objects.all().prefetch_related(
        "lideres", "miembros", "bandas_aliadas", "bandas_rivales"
    )
    if filtros["nombre"]:
        bandas_queryset = bandas_queryset.filter(nombres__icontains=filtros["nombre"])
    if filtros["zona"]:
        bandas_queryset = bandas_queryset.filter(
            zonas_influencia__icontains=filtros["zona"]
        )
    if filtros["aliada"]:
        bandas_queryset = bandas_queryset.filter(
            bandas_aliadas__nombres__icontains=filtros["aliada"]
        )
    if filtros["rival"]:
        bandas_queryset = bandas_queryset.filter(
            bandas_rivales__nombres__icontains=filtros["rival"]
        )
    bandas_queryset = bandas_queryset.distinct()
    bandas = sorted(
        bandas_queryset,
        key=lambda banda: (banda.nombre_principal or "").lower(),
    )

    return render(
        request,
        "bandas_criminales.html",
        {
            "form": form,
            "bandas": bandas,
            "modo_edicion": True,
            "banda_actual": banda,
            "form_action": reverse("banda_criminal_editar", args=[banda.pk]),
            "filtros": filtros,
        },
    )


@login_required
def banda_criminal_eliminar_view(request, pk):
    banda = get_object_or_404(BandaCriminal, pk=pk)

    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        nombre = banda.nombres_como_texto or "Sin nombre"
        banda.delete()
        log_actividad(
            request,
            TipoActividad.OTRO,
            f"Banda criminal eliminada (ID {pk}, {nombre}).",
        )
        messages.success(request, "Banda criminal eliminada correctamente.")
        return redirect("bandas_criminales")

    messages.error(request, "Método no permitido para eliminar bandas criminales.")
    return redirect("bandas_criminales")


@login_required
def articulo_editar_view(request, id):
    articulo = get_object_or_404(
        Articulo.objects.select_related("generado_por", "categoria").prefetch_related(
            "solicitudes_info__respondido_por"
        ),
        pk=id,
    )

    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    if (
        request.user.userprofile.rol == Roles.REDACCION
        and articulo.generado_por != request.user
    ):
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = ArticuloForm(request.POST, instance=articulo)
        if form.is_valid():
            form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Artículo actualizado desde panel: {articulo.titulo} (ID {articulo.pk}).",
            )
            messages.success(request, "Artículo actualizado correctamente.")
            return redirect("articulo_editar", id=articulo.id)
    else:
        form = ArticuloForm(instance=articulo)

    solicitudes_relacionadas = (
        articulo.solicitudes_info.select_related("respondido_por")
        .all()
        .order_by("-fecha_creacion")
    )

    return render(
        request,
        "articulo_editar.html",
        {
            "articulo": articulo,
            "form": form,
            "solicitudes_relacionadas": solicitudes_relacionadas,
        },
    )


@login_required
def osint_panel_view(request):
    if request.user.userprofile.rol not in [Roles.PRENSA, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    herramientas = HerramientaOSINT.objects.all()
    return render(
        request,
        "osint_panel.html",
        {
            "herramientas": herramientas,
        },
    )


@login_required
def solicitud_info_portal_view(request):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitudes_usuario = (
        SolicitudInfo.objects.select_related("respondido_por")
        .prefetch_related("articulos")
        .filter(usuario_creador=request.user)
        .order_by("-fecha_creacion")
    )

    paginator = Paginator(solicitudes_usuario, 10)
    page_number = request.GET.get("page")
    solicitudes_page = paginator.get_page(page_number)

    if request.method == "POST":
        form = SolicitudInfoForm(request.POST, user=request.user)
        if form.is_valid():
            solicitud = form.save(commit=False)
            solicitud.usuario_creador = request.user
            solicitud.save()
            articulos_relacionados = form.cleaned_data.get("articulos")
            if articulos_relacionados:
                solicitud.articulos.set(articulos_relacionados)
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Solicitud de información creada (# {solicitud.pk}).",
            )
            messages.success(request, "Solicitud enviada correctamente.")
            return redirect("solicitud_info_portal")
        messages.error(
            request,
            "No se pudo crear la solicitud. Por favor, revisá los campos destacados.",
        )
    else:
        form = SolicitudInfoForm(user=request.user)

    articulos_queryset = form.fields["articulos"].queryset.select_related("categoria")
    categorias_articulos = (
        articulos_queryset.exclude(categoria__isnull=True)
        .values_list("categoria__nombre", flat=True)
        .distinct()
        .order_by("categoria__nombre")
    )

    return render(
        request,
        "solicitud_info_portal.html",
        {
            "form": form,
            "solicitudes_usuario": solicitudes_usuario,
            "solicitudes_page": solicitudes_page,
            "articulos_disponibles": articulos_queryset,
            "categorias_articulos": list(categorias_articulos),
            "estados_articulo": Articulo.Estado.choices,
            "modo_edicion": False,
            "solicitud_en_edicion": None,
            "form_action": "",
        },
    )


@login_required
def solicitud_info_portal_detalle_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitud = get_object_or_404(
        SolicitudInfo.objects.select_related("usuario_creador", "respondido_por").prefetch_related("articulos"),
        pk=pk,
    )

    if request.user.userprofile.rol == Roles.REDACCION and solicitud.usuario_creador != request.user:
        return render(request, "403.html", status=403)

    return render(
        request,
        "solicitud_info_portal_detalle.html",
        {
            "solicitud": solicitud,
        },
    )


# NUEVA VISTA: Editar solicitud de información
@login_required
def solicitud_info_portal_editar_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitud = get_object_or_404(
        SolicitudInfo.objects.select_related("usuario_creador", "respondido_por").prefetch_related("articulos"),
        pk=pk,
    )

    if request.user.userprofile.rol == Roles.REDACCION and solicitud.usuario_creador != request.user:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = SolicitudInfoForm(request.POST, instance=solicitud, user=request.user)
        if form.is_valid():
            solicitud = form.save(commit=False)
            solicitud.save()
            articulos_relacionados = form.cleaned_data.get("articulos")
            if articulos_relacionados is not None:
                solicitud.articulos.set(articulos_relacionados)
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Solicitud de información editada (# {solicitud.pk}).",
            )
            messages.success(request, "Solicitud actualizada correctamente.")
            return redirect("solicitud_info_portal")
        messages.error(
            request,
            "No se pudo actualizar la solicitud. Por favor, revisá los campos destacados.",
        )
    else:
        form = SolicitudInfoForm(instance=solicitud, user=request.user)

    articulos_queryset = form.fields["articulos"].queryset.select_related("categoria")
    categorias_articulos = (
        articulos_queryset.exclude(categoria__isnull=True)
        .values_list("categoria__nombre", flat=True)
        .distinct()
        .order_by("categoria__nombre")
    )

    return render(
        request,
        "solicitud_info_portal_editar.html",
        {
            "form": form,
            "solicitud": solicitud,
            "articulos_disponibles": articulos_queryset,
            "categorias_articulos": list(categorias_articulos),
            "estados_articulo": Articulo.Estado.choices,
        },
    )


@login_required
def solicitud_info_portal_eliminar_view(request, pk):
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitud = get_object_or_404(SolicitudInfo, pk=pk)

    if request.user.userprofile.rol == Roles.REDACCION and solicitud.usuario_creador != request.user:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        solicitud_id = solicitud.pk
        solicitud.delete()
        log_actividad(
            request,
            TipoActividad.OTRO,
            f"Solicitud de información eliminada (# {solicitud_id}).",
        )
        messages.success(request, "Solicitud eliminada correctamente.")
    else:
        messages.error(request, "Operación no válida.")
    return redirect("solicitud_info_portal")


@login_required
def solicitud_info_crear_view(request):
    """
    Mantener compatibilidad con rutas antiguas redirigiendo al nuevo portal.
    """
    return redirect("solicitud_info_portal")


@login_required
def solicitudes_info_list_view(request):
    if request.user.userprofile.rol not in [Roles.PRENSA, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitudes = (
        SolicitudInfo.objects.select_related("usuario_creador", "respondido_por")
        .order_by("-fecha_creacion")
    )

    pendientes = solicitudes.filter(fecha_respuesta__isnull=True).count()

    return render(
        request,
        "solicitudes_info_list.html",
        {
            "solicitudes": solicitudes,
            "pendientes": pendientes,
        },
    )


@login_required
def solicitud_info_detalle_view(request, pk):
    if request.user.userprofile.rol not in [Roles.PRENSA, Roles.ADMIN]:
        return render(request, "403.html", status=403)

    solicitud = get_object_or_404(
        SolicitudInfo.objects.select_related("usuario_creador", "respondido_por"), pk=pk
    )

    if solicitud.fecha_lectura is None:
        solicitud.fecha_lectura = now()
        solicitud.save(update_fields=["fecha_lectura"])

    if request.method == "POST":
        form = SolicitudInfoRespuestaForm(request.POST, instance=solicitud)
        if form.is_valid():
            solicitud = form.save(commit=False)
            solicitud.respondido_por = request.user
            if solicitud.fecha_lectura is None:
                solicitud.fecha_lectura = now()
            solicitud.fecha_respuesta = now()
            solicitud.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Solicitud de información #{solicitud.pk} respondida.",
            )
            messages.success(request, "Respuesta registrada correctamente.")
            return redirect("solicitud_info_detalle", pk=solicitud.pk)
        messages.error(request, "No se pudo registrar la respuesta. Por favor, revisá el formulario.")
    else:
        form = SolicitudInfoRespuestaForm(instance=solicitud)

    return render(
        request,
        "solicitud_info_detalle.html",
        {
            "solicitud": solicitud,
            "respuesta_form": form,
        },
    )


@login_required
def clasificacion_view(request):
    if request.user.userprofile.rol in [Roles.CLASIFICACION, Roles.ADMIN]:
        categorias = Categoria.objects.all()
        return render(request, "clasificacion.html", {"categorias": categorias})
    return render(request, "403.html", status=403)


@login_required
def actividad_view(request):
    if request.user.userprofile.rol == Roles.ADMIN:
        return render(request, "actividad.html")
    return render(request, "403.html", status=403)


@login_required
def actividad_debug_view(request):
    # Acceso a esta vista sin restricción de rol, asumiendo que el login_required es suficiente para el debugging.
    return render(request, "actividad_debug.html")


# -------------------------- EXPORTACIONES --------------------------


def exportar_actividades_excel(request):
    actividades = Actividad.objects.select_related("usuario").order_by("-fecha_hora")
    usuario = request.GET.get("usuario")
    tipo = request.GET.get("tipo")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")

    if usuario:
        actividades = actividades.filter(usuario__username=usuario)
    if tipo:
        actividades = actividades.filter(tipo=tipo)
    if desde:
        actividades = actividades.filter(fecha_hora__date__gte=parse_date(desde))
    if hasta:
        actividades = actividades.filter(fecha_hora__date__lte=parse_date(hasta))

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Actividades"
    ws.append(["Fecha", "Usuario", "Tipo", "Descripción"])

    for act in actividades:
        ws.append(
            [
                act.fecha_hora.strftime("%Y-%m-%d %H:%M:%S"),
                act.usuario.username if act.usuario else "(anónimo)",
                act.get_tipo_display(),
                act.descripcion,
            ]
        )

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = "attachment; filename=actividades.xlsx"
    wb.save(response)
    return response


@login_required
def exportar_actividades_pdf(request):
    actividades = Actividad.objects.select_related("usuario").order_by("-fecha_hora")
    usuario = request.GET.get("usuario")
    tipo = request.GET.get("tipo")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")

    if usuario:
        actividades = actividades.filter(usuario__username=usuario)
    if tipo:
        actividades = actividades.filter(tipo=tipo)
    if desde:
        actividades = actividades.filter(fecha_hora__date__gte=parse_date(desde))
    if hasta:
        actividades = actividades.filter(fecha_hora__date__lte=parse_date(hasta))

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=actividades.pdf"

    doc = SimpleDocTemplate(response, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    elements = []
    styles = getSampleStyleSheet()

    # Título y filtros
    elements.append(Paragraph("<b>Registro de Actividades</b>", styles['Heading1']))
    
    filtros = []
    if usuario:
        filtros.append(f"Usuario: <b>{usuario}</b>")
    if tipo:
        # Aquí, 'tipo' es el valor de la opción, no el display. Podrías buscar el display si fuera necesario.
        filtros.append(f"Tipo: <b>{tipo}</b>") 
    if desde:
        filtros.append(f"Desde: <b>{desde}</b>")
    if hasta:
        filtros.append(f"Hasta: <b>{hasta}</b>")

    if filtros:
        filtros_str = " | ".join(filtros)
        elements.append(Paragraph(f"Filtros aplicados: {filtros_str}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Tabla
    data = [["Fecha", "Usuario", "Tipo", "Descripción"]]

    for act in actividades:
        # Usa el estilo 'Normal' para que el texto de la descripción se ajuste a la celda
        data.append([
            act.fecha_hora.strftime('%Y-%m-%d %H:%M:%S'),
            act.usuario.username if act.usuario else '(anónimo)',
            act.get_tipo_display(),
            Paragraph(act.descripcion, styles['Normal'])
        ])

    table = Table(data, colWidths=[4*cm, 4*cm, 3*cm, 7.5*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, -1), 8), # Reducir fuente para que quepa más
    ]))

    elements.append(table)
    doc.build(elements)

    return response


@login_required
def exportar_articulo_pdf(request, id):
    if HTML is None:
        return _weasyprint_unavailable_response(request, "redaccion")
    articulo = get_object_or_404(
        Articulo.objects.select_related("generado_por", "categoria").prefetch_related(
            "links_incluidos"
        ),
        pk=id,
    )

    context = {
        "titulo": articulo.titulo,
        "descripcion": articulo.descripcion,
        "categoria": (
            articulo.categoria.nombre if articulo.categoria else "Sin categoría"
        ),
        "fecha": articulo.fecha_creacion,
        "autor": (
            articulo.generado_por.username if articulo.generado_por else "Desconocido"
        ),
        "links": articulo.links_incluidos.all(),
    }

    # Asumiendo que "articulo_pdf.html" existe y es válido
    html = render_to_string("articulo_pdf.html", context)
    
    # Manejo de la URL base para recursos (CSS, imágenes)
    pdf_file = HTML(string=html, base_url=request.build_absolute_uri()).write_pdf()

    # Armado del nombre dinámico
    fecha_str = localtime(articulo.fecha_creacion).strftime("%Y-%m-%d")
    titulo_slug = slugify(articulo.titulo)[:50] or "articulo"
    nombre_archivo = f"Articulo_{titulo_slug}_{fecha_str}.pdf"

    response = HttpResponse(pdf_file, content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="{nombre_archivo}"'
    return response


@csrf_exempt
def registrar_clic_link(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            url = data.get("url", "URL desconocida")
            log_actividad(request, TipoActividad.CLIC_LINK, f"Click en link: {url}")
            return JsonResponse({"status": "ok"})
        except json.JSONDecodeError:
            return JsonResponse({"status": "error", "mensaje": "Datos JSON inválidos"}, status=400)
        except Exception as e:
            logger.error(f"Error al registrar click: {e}")
            return JsonResponse({"status": "error", "mensaje": str(e)}, status=400)
    return JsonResponse({"status": "método no permitido"}, status=405)


# -------------------------- API VIEWSETS --------------------------


class UserViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Endpoint para listar usuarios con perfil.
    """

    queryset = User.objects.all().select_related("userprofile")
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]


class UserProfileViewSet(viewsets.ModelViewSet):
    queryset = UserProfile.objects.select_related("user")
    serializer_class = UserProfileSerializer
    permission_classes = [IsAuthenticated]


class DiarioDigitalViewSet(viewsets.ModelViewSet):
    queryset = DiarioDigital.objects.all()
    serializer_class = DiarioDigitalSerializer
    permission_classes = [IsAuthenticated]


class RedSocialViewSet(viewsets.ModelViewSet):
    queryset = RedSocial.objects.all()
    serializer_class = RedSocialSerializer
    permission_classes = [IsAuthenticated]


class TvDigitalViewSet(viewsets.ModelViewSet):
    queryset = TvDigital.objects.all()
    serializer_class = TvDigitalSerializer
    permission_classes = [IsAuthenticated]


class RadioDigitalViewSet(viewsets.ModelViewSet):
    queryset = RadioDigital.objects.all()
    serializer_class = RadioDigitalSerializer
    permission_classes = [IsAuthenticated]


class CategoriaViewSet(viewsets.ModelViewSet):
    queryset = Categoria.objects.all()
    serializer_class = CategoriaSerializer
    permission_classes = [IsAuthenticated]


class LinkRelevanteViewSet(viewsets.ModelViewSet):
    queryset = (
        LinkRelevante.objects.select_related("cargado_por", "diario_digital")
        .prefetch_related("categorias")
        .order_by("-fecha_carga")
    )
    permission_classes = [IsAuthenticated]
    serializer_class = LinkRelevanteSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        # Si el usuario tiene el rol prensa, solo ve sus propios links
        if hasattr(user, "userprofile"):
            rol = user.userprofile.rol
            if rol == Roles.PRENSA:
                queryset = queryset.filter(cargado_por=user)
            # Los roles CLIENTE e INFORMES no deberían usar esta API a menos que sea necesario
            # para ver links aprobados. 
            elif rol in [Roles.CLIENTE, Roles.INFORMES]:
                queryset = queryset.filter(estado='aprobado') 
            elif rol in [Roles.CLASIFICACION, Roles.REDACCION, Roles.GERENCIA, Roles.ADMIN]:
                pass  # pueden ver todos los links

        # Filtramos por los parámetros GET si vienen
        estado = self.request.query_params.get("estado")
        fecha_inicio = self.request.query_params.get("fecha_inicio")
        fecha_fin = self.request.query_params.get("fecha_fin")
        categoria_id = self.request.query_params.get("categoria_id")

        if estado:
            queryset = queryset.filter(estado=estado)
        if fecha_inicio:
            queryset = queryset.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
        if fecha_fin:
            queryset = queryset.filter(fecha_carga__date__lte=parse_date(fecha_fin))
        if categoria_id:
            queryset = queryset.filter(categorias__id=categoria_id)

        return queryset.distinct()

    def _resolver_fuentes(self, url):
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower().replace("www.", "")

        diario = None
        for d in DiarioDigital.objects.all():
            diario_domain = urlparse(d.url_principal).netloc.lower().replace("www.", "")
            if diario_domain == domain or domain.endswith(f".{diario_domain}"):
                diario = d
                break

        return diario

    def perform_create(self, serializer):
        url = serializer.validated_data.get("url", "")
        diario = self._resolver_fuentes(url)

        link = serializer.save(
            cargado_por=self.request.user,
            diario_digital=diario,
        )
        log_actividad(
            self.request,
            TipoActividad.CARGA_LINK,
            "Se cargó un nuevo link: {url} (Fuente asignada: {fuente})".format(
                url=link.url,
                fuente=(
                    diario.nombre
                    if diario
                    else "Ninguna"
                ),
            ),
        )

    def perform_update(self, serializer):
        instance = self.get_object()
        original_estado = instance.estado
        original_categorias_ids = set(instance.categorias.values_list("id", flat=True))

        link_instance = serializer.save()

        if "url" in serializer.validated_data and "diario_digital" not in serializer.validated_data:
            diario = self._resolver_fuentes(link_instance.url)
            update_required = False
            update_fields = []
            if link_instance.diario_digital != diario:
                link_instance.diario_digital = diario
                update_required = True
                update_fields.append("diario_digital")
            if update_required:
                link_instance.save(update_fields=update_fields)

        # Log de cambio de estado
        if original_estado != link_instance.estado:
            log_actividad(
                self.request,
                TipoActividad.CAMBIO_ESTADO,
                f"Se cambió el estado del link ID {link_instance.id} de '{original_estado}' a '{link_instance.estado}'. URL: {link_instance.url}",
            )
            
        # Actualizar fecha de aprobación si pasa a 'aprobado'
        if original_estado != 'aprobado' and link_instance.estado == 'aprobado':
            link_instance.fecha_aprobacion = datetime.now()
            link_instance.save(update_fields=['fecha_aprobacion'])


        current_categorias_ids = set(
            link_instance.categorias.values_list("id", flat=True)
        )
        # Log de cambio de clasificación
        if original_categorias_ids != current_categorias_ids:
            nuevas_categorias_nombres = ", ".join(
                [c.nombre for c in link_instance.categorias.all()]
            )
            log_actividad(
                self.request,
                TipoActividad.CLASIFICACION_LINK,
                f"Se actualizaron las categorías del link ID {link_instance.id}. Nuevas categorías: [{nuevas_categorias_nombres}]. URL: {link_instance.url}",
            )

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            logger.error("ERRORES DEL SERIALIZER:", serializer.errors)
            return Response(serializer.errors, status=400)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=201, headers=headers)


class LinkRedSocialViewSet(viewsets.ModelViewSet):
    queryset = (
        LinkRedSocial.objects.select_related("cargado_por", "red_social")
        .prefetch_related("categorias")
        .order_by("-fecha_carga")
    )
    permission_classes = [IsAuthenticated]
    serializer_class = LinkRedSocialSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if hasattr(user, "userprofile"):
            rol = user.userprofile.rol
            if rol == Roles.PRENSA:
                queryset = queryset.filter(cargado_por=user)
            elif rol in [Roles.CLIENTE, Roles.INFORMES]:
                queryset = queryset.filter(estado=EstadoLink.APROBADO)

        estado = self.request.query_params.get("estado")
        fecha_inicio = self.request.query_params.get("fecha_inicio")
        fecha_fin = self.request.query_params.get("fecha_fin")
        categoria_id = self.request.query_params.get("categoria_id")
        red_social_id = self.request.query_params.get("red_social_id")
        solo_propios = self.request.query_params.get("solo_propios")

        if estado:
            queryset = queryset.filter(estado=estado)
        if fecha_inicio:
            queryset = queryset.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
        if fecha_fin:
            queryset = queryset.filter(fecha_carga__date__lte=parse_date(fecha_fin))
        if categoria_id:
            queryset = queryset.filter(categorias__id=categoria_id)
        if red_social_id:
            queryset = queryset.filter(red_social__id=red_social_id)
        if solo_propios:
            queryset = queryset.filter(cargado_por=user)

        return queryset.distinct()

    def _resolver_red_social(self, url):
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower().replace("www.", "")

        for red in RedSocial.objects.all():
            red_domain = urlparse(red.url_principal).netloc.lower().replace("www.", "")
            if red_domain == domain or domain.endswith(f".{red_domain}"):
                return red
        return None

    def perform_create(self, serializer):
        url = serializer.validated_data.get("url", "")
        red_social = serializer.validated_data.get("red_social") or self._resolver_red_social(url)

        link = serializer.save(
            cargado_por=self.request.user,
            red_social=red_social,
        )
        log_actividad(
            self.request,
            TipoActividad.CARGA_LINK,
            "Se cargó un nuevo link de red social: {url} (Red asignada: {red})".format(
                url=link.url,
                red=link.red_social.nombre if link.red_social else "Ninguna",
            ),
        )

    def perform_update(self, serializer):
        instance = self.get_object()
        original_estado = instance.estado
        original_categorias_ids = set(instance.categorias.values_list("id", flat=True))

        link_instance = serializer.save()

        if "url" in serializer.validated_data and "red_social" not in serializer.validated_data:
            red_social = self._resolver_red_social(link_instance.url)
            if link_instance.red_social != red_social:
                link_instance.red_social = red_social
                link_instance.save(update_fields=["red_social"])

        if original_estado != link_instance.estado:
            log_actividad(
                self.request,
                TipoActividad.CAMBIO_ESTADO,
                f"Se cambió el estado del link de red social ID {link_instance.id} de '{original_estado}' a '{link_instance.estado}'. URL: {link_instance.url}",
            )

        if original_estado != EstadoLink.APROBADO and link_instance.estado == EstadoLink.APROBADO:
            link_instance.fecha_aprobacion = datetime.now()
            link_instance.save(update_fields=["fecha_aprobacion"])

        current_categorias_ids = set(
            link_instance.categorias.values_list("id", flat=True)
        )
        if original_categorias_ids != current_categorias_ids:
            nuevas_categorias_nombres = ", ".join(
                [c.nombre for c in link_instance.categorias.all()]
            )
            log_actividad(
                self.request,
                TipoActividad.CLASIFICACION_LINK,
                f"Se actualizaron las categorías del link de red social ID {link_instance.id}. Nuevas categorías: [{nuevas_categorias_nombres}]. URL: {link_instance.url}",
            )

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            logger.error("ERRORES DEL SERIALIZER (LinkRedSocial): %s", serializer.errors)
            return Response(serializer.errors, status=400)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=201, headers=headers)


class LinkTvDigitalViewSet(viewsets.ModelViewSet):
    queryset = (
        LinkTvDigital.objects.select_related("cargado_por", "tv_digital")
        .prefetch_related("categorias")
        .order_by("-fecha_carga")
    )
    permission_classes = [IsAuthenticated]
    serializer_class = LinkTvDigitalSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if hasattr(user, "userprofile"):
            rol = user.userprofile.rol
            if rol == Roles.PRENSA:
                queryset = queryset.filter(cargado_por=user)
            elif rol in [Roles.CLIENTE, Roles.INFORMES]:
                queryset = queryset.filter(estado=EstadoLink.APROBADO)

        estado = self.request.query_params.get("estado")
        fecha_inicio = self.request.query_params.get("fecha_inicio")
        fecha_fin = self.request.query_params.get("fecha_fin")
        categoria_id = self.request.query_params.get("categoria_id")
        tv_digital_id = self.request.query_params.get("tv_digital_id")
        solo_propios = self.request.query_params.get("solo_propios")

        if estado:
            queryset = queryset.filter(estado=estado)
        if fecha_inicio:
            queryset = queryset.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
        if fecha_fin:
            queryset = queryset.filter(fecha_carga__date__lte=parse_date(fecha_fin))
        if categoria_id:
            queryset = queryset.filter(categorias__id=categoria_id)
        if tv_digital_id:
            queryset = queryset.filter(tv_digital__id=tv_digital_id)
        if solo_propios:
            queryset = queryset.filter(cargado_por=user)

        return queryset.distinct()

    def _resolver_tv_digital(self, url):
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower().replace("www.", "")

        for tv in TvDigital.objects.all():
            tv_domain = urlparse(tv.url_principal).netloc.lower().replace("www.", "")
            if tv_domain == domain or domain.endswith(f".{tv_domain}"):
                return tv
        return None

    def perform_create(self, serializer):
        url = serializer.validated_data.get("url", "")
        tv_digital = serializer.validated_data.get("tv_digital") or self._resolver_tv_digital(url)

        link = serializer.save(
            cargado_por=self.request.user,
            tv_digital=tv_digital,
        )
        log_actividad(
            self.request,
            TipoActividad.CARGA_LINK,
            "Se cargó un nuevo link de TV digital: {url} (Fuente asignada: {tv})".format(
                url=link.url,
                tv=link.tv_digital.nombre if link.tv_digital else "Ninguna",
            ),
        )

    def perform_update(self, serializer):
        instance = self.get_object()
        original_estado = instance.estado
        original_categorias_ids = set(instance.categorias.values_list("id", flat=True))

        link_instance = serializer.save()

        if "url" in serializer.validated_data and "tv_digital" not in serializer.validated_data:
            tv_digital = self._resolver_tv_digital(link_instance.url)
            if link_instance.tv_digital != tv_digital:
                link_instance.tv_digital = tv_digital
                link_instance.save(update_fields=["tv_digital"])

        if original_estado != link_instance.estado:
            log_actividad(
                self.request,
                TipoActividad.CAMBIO_ESTADO,
                f"Se cambió el estado del link de TV digital ID {link_instance.id} de '{original_estado}' a '{link_instance.estado}'. URL: {link_instance.url}",
            )

        if original_estado != EstadoLink.APROBADO and link_instance.estado == EstadoLink.APROBADO:
            link_instance.fecha_aprobacion = datetime.now()
            link_instance.save(update_fields=["fecha_aprobacion"])

        current_categorias_ids = set(
            link_instance.categorias.values_list("id", flat=True)
        )
        if original_categorias_ids != current_categorias_ids:
            nuevas_categorias_nombres = ", ".join(
                [c.nombre for c in link_instance.categorias.all()]
            )
            log_actividad(
                self.request,
                TipoActividad.CLASIFICACION_LINK,
                f"Se actualizaron las categorías del link de TV digital ID {link_instance.id}. Nuevas categorías: [{nuevas_categorias_nombres}]. URL: {link_instance.url}",
            )

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            logger.error("ERRORES DEL SERIALIZER (LinkTvDigital): %s", serializer.errors)
            return Response(serializer.errors, status=400)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=201, headers=headers)


class LinkRadioDigitalViewSet(viewsets.ModelViewSet):
    queryset = (
        LinkRadioDigital.objects.select_related("cargado_por", "radio_digital")
        .prefetch_related("categorias")
        .order_by("-fecha_carga")
    )
    permission_classes = [IsAuthenticated]
    serializer_class = LinkRadioDigitalSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if hasattr(user, "userprofile"):
            rol = user.userprofile.rol
            if rol == Roles.PRENSA:
                queryset = queryset.filter(cargado_por=user)
            elif rol in [Roles.CLIENTE, Roles.INFORMES]:
                queryset = queryset.filter(estado=EstadoLink.APROBADO)

        estado = self.request.query_params.get("estado")
        fecha_inicio = self.request.query_params.get("fecha_inicio")
        fecha_fin = self.request.query_params.get("fecha_fin")
        categoria_id = self.request.query_params.get("categoria_id")
        radio_digital_id = self.request.query_params.get("radio_digital_id")
        solo_propios = self.request.query_params.get("solo_propios")

        if estado:
            queryset = queryset.filter(estado=estado)
        if fecha_inicio:
            queryset = queryset.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
        if fecha_fin:
            queryset = queryset.filter(fecha_carga__date__lte=parse_date(fecha_fin))
        if categoria_id:
            queryset = queryset.filter(categorias__id=categoria_id)
        if radio_digital_id:
            queryset = queryset.filter(radio_digital__id=radio_digital_id)
        if solo_propios:
            queryset = queryset.filter(cargado_por=user)

        return queryset.distinct()

    def _resolver_radio_digital(self, url):
        parsed_url = urlparse(url)
        domain = parsed_url.netloc.lower().replace("www.", "")

        for radio in RadioDigital.objects.all():
            radio_domain = urlparse(radio.url_principal).netloc.lower().replace("www.", "")
            if radio_domain == domain or domain.endswith(f".{radio_domain}"):
                return radio
        return None

    def perform_create(self, serializer):
        url = serializer.validated_data.get("url", "")
        radio_digital = serializer.validated_data.get("radio_digital") or self._resolver_radio_digital(url)

        link = serializer.save(
            cargado_por=self.request.user,
            radio_digital=radio_digital,
        )
        log_actividad(
            self.request,
            TipoActividad.CARGA_LINK,
            "Se cargó un nuevo link de Radio digital: {url} (Fuente asignada: {radio})".format(
                url=link.url,
                radio=link.radio_digital.nombre if link.radio_digital else "Ninguna",
            ),
        )

    def perform_update(self, serializer):
        instance = self.get_object()
        original_estado = instance.estado
        original_categorias_ids = set(instance.categorias.values_list("id", flat=True))

        link_instance = serializer.save()

        if "url" in serializer.validated_data and "radio_digital" not in serializer.validated_data:
            radio_digital = self._resolver_radio_digital(link_instance.url)
            if link_instance.radio_digital != radio_digital:
                link_instance.radio_digital = radio_digital
                link_instance.save(update_fields=["radio_digital"])

        if original_estado != link_instance.estado:
            log_actividad(
                self.request,
                TipoActividad.CAMBIO_ESTADO,
                f"Se cambió el estado del link de Radio digital ID {link_instance.id} de '{original_estado}' a '{link_instance.estado}'. URL: {link_instance.url}",
            )

        if original_estado != EstadoLink.APROBADO and link_instance.estado == EstadoLink.APROBADO:
            link_instance.fecha_aprobacion = datetime.now()
            link_instance.save(update_fields=["fecha_aprobacion"])

        current_categorias_ids = set(
            link_instance.categorias.values_list("id", flat=True)
        )
        if original_categorias_ids != current_categorias_ids:
            nuevas_categorias_nombres = ", ".join(
                [c.nombre for c in link_instance.categorias.all()]
            )
            log_actividad(
                self.request,
                TipoActividad.CLASIFICACION_LINK,
                f"Se actualizaron las categorías del link de Radio digital ID {link_instance.id}. Nuevas categorías: [{nuevas_categorias_nombres}]. URL: {link_instance.url}",
            )

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            logger.error("ERRORES DEL SERIALIZER (LinkRadioDigital): %s", serializer.errors)
            return Response(serializer.errors, status=400)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=201, headers=headers)


class ArticuloViewSet(viewsets.ModelViewSet):
    queryset = (
        Articulo.objects.all()
        .select_related("categoria", "generado_por")
        .prefetch_related("links_incluidos", "hechos_destacados", "solicitudes_info")
    )
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        # CORRECCIÓN: Usar ArticuloSerializer para todas las acciones por ahora
        # para evitar problemas con ArticuloDetailSerializer
        return ArticuloSerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        user = self.request.user

        rol_usuario = getattr(getattr(user, "userprofile", None), "rol", None)

        if rol_usuario != Roles.ADMIN:
            queryset = queryset.filter(generado_por=user)

        # Filtros adicionales
        fecha_inicio = self.request.query_params.get("fecha_inicio")
        fecha_fin = self.request.query_params.get("fecha_fin")
        categoria_id = self.request.query_params.get("categoria_id")

        if fecha_inicio:
            queryset = queryset.filter(
                fecha_creacion__date__gte=parse_date(fecha_inicio)
            )
        if fecha_fin:
            queryset = queryset.filter(
                fecha_creacion__date__lte=parse_date(fecha_fin)
            )
        if categoria_id:
            queryset = queryset.filter(categoria__id=categoria_id)

        return queryset

    def perform_create(self, serializer):
        articulo = serializer.save(generado_por=self.request.user)
        log_actividad(self.request, TipoActividad.CARGA_ARTICULO, f"Se cargó un nuevo artículo: {articulo.titulo}")

    @action(detail=True, methods=["get"], url_path="consultar")
    def consultar(self, request, pk=None):
        """
        Permite consultar el artículo, actualizar revisado_redactor y devolverlo en detalle.
        """
        try:
            # Se usa el queryset base para asegurar los prefetches
            articulo = self.get_queryset().get(pk=pk)
        except Articulo.DoesNotExist:
            return Response(
                {"status": "error", "msg": "Artículo no encontrado."}, status=404
            )

        # Marcar como revisado solo si el usuario no es quien lo generó (para evitar auto-revisión)
        if articulo.generado_por != request.user:
            articulo.revisado_redactor = True
            articulo.save(update_fields=["revisado_redactor"])
            log_actividad(request, TipoActividad.OTRO, f"Artículo ID {pk} marcado como revisado por redactor.")

        # Serializar con el Detail
        serializer = ArticuloDetailSerializer(articulo)
        return Response({"status": "ok", "articulo": serializer.data})


class ActividadPagination(PageNumberPagination):
    page_size = 20
    page_size_query_param = "page_size"
    max_page_size = 100


class ActividadViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Actividad.objects.select_related("usuario").order_by("-fecha_hora")
    serializer_class = ActividadSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = ActividadPagination

    def get_queryset(self):
        qs = super().get_queryset()
        usuario = self.request.query_params.get("usuario")
        tipo = self.request.query_params.get("tipo")
        desde = self.request.query_params.get("desde")
        hasta = self.request.query_params.get("hasta")
        if usuario:
            qs = qs.filter(usuario__username=usuario)
        if tipo:
            qs = qs.filter(tipo=tipo)
        if desde:
            qs = qs.filter(fecha_hora__date__gte=parse_date(desde))
        if hasta:
            qs = qs.filter(fecha_hora__date__lte=parse_date(hasta))
        return qs


class InformeIndividualViewSet(viewsets.ModelViewSet):
    queryset = InformeIndividual.objects.prefetch_related(
        "articulos", "vehiculos", "empleadores", "alias", "vinculos", "domicilio", "telefono"
    )
    serializer_class = InformeIndividualSerializer
    permission_classes = [IsAuthenticated]


# Optional CRUD for auxiliary models
class DomicilioViewSet(viewsets.ModelViewSet):
    queryset = Domicilio.objects.all()
    serializer_class = DomicilioSerializer
    permission_classes = [IsAuthenticated]


class VehiculoViewSet(viewsets.ModelViewSet):
    queryset = Vehiculo.objects.all()
    serializer_class = VehiculoSerializer
    permission_classes = [IsAuthenticated]


class EmpleadorViewSet(viewsets.ModelViewSet):
    queryset = Empleador.objects.all()
    serializer_class = EmpleadorSerializer
    permission_classes = [IsAuthenticated]


class AliasViewSet(viewsets.ModelViewSet):
    queryset = Alias.objects.all()
    serializer_class = AliasSerializer
    permission_classes = [IsAuthenticated]


class VinculoViewSet(viewsets.ModelViewSet):
    queryset = Vinculo.objects.all()
    serializer_class = VinculoSerializer
    permission_classes = [IsAuthenticated]


class HechoDestacadoViewSet(viewsets.ModelViewSet):
    queryset = HechoDestacado.objects.all()
    serializer_class = HechoDestacadoSerializer
    permission_classes = [IsAuthenticated]


# -------------------- FUNCTION-BASED API PARA PANEL DE CLASIFICACIÓN --------------------

@csrf_exempt
def api_links_list(request):
    try:
        # Obtener parámetros de filtro
        fecha_inicio = request.GET.get('fecha_inicio')
        fecha_fin = request.GET.get('fecha_fin')
        diario_id = request.GET.get('diario_id')
        estado = request.GET.get('estado')
        categoria_id = request.GET.get('categoria_id')
        solo_propios = request.GET.get('solo_propios')
        fuente = request.GET.get('fuente')
        red_social_id = request.GET.get('red_social_id')
        tv_digital_id = request.GET.get('tv_digital_id')
        radio_digital_id = request.GET.get('radio_digital_id')

        data = []

        if fuente == 'red_social':
            links = LinkRedSocial.objects.select_related('red_social').prefetch_related('categorias').all()

            if fecha_inicio:
                links = links.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
            if fecha_fin:
                links = links.filter(fecha_carga__date__lte=parse_date(fecha_fin))
            if estado:
                if estado.startswith('estado!='):
                    exclude_estado = estado.split('!=')[1]
                    links = links.exclude(estado=exclude_estado)
                else:
                    links = links.filter(estado=estado)
            if categoria_id:
                links = links.filter(categorias__id=categoria_id)
            if red_social_id:
                links = links.filter(red_social__id=red_social_id)
            if solo_propios:
                links = links.filter(cargado_por=request.user)

            for link in links:
                data.append({
                    'id': link.id,
                    'url': link.url,
                    'estado': getattr(link, 'estado', EstadoLink.PENDIENTE),
                    'fecha_carga': link.fecha_carga.isoformat() if link.fecha_carga else None,
                    'revisado_clasificador': getattr(link, 'revisado_clasificador', False),
                    'diario_logo_url': None,
                    'diario_nombre': None,
                    'diario_digital': None,
                    'red_social': link.red_social.id if link.red_social else None,
                    'red_social_nombre': link.red_social.nombre if link.red_social else 'Sin asignar',
                    'red_social_logo_url': link.red_social.logo.url if link.red_social and getattr(link.red_social, 'logo', None) else None,
                    'tv_digital': None,
                    'tv_digital_nombre': None,
                    'tv_digital_logo_url': None,
                    'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
                })
        elif fuente == 'tv_digital':
            links = LinkTvDigital.objects.select_related('tv_digital').prefetch_related('categorias').all()

            if fecha_inicio:
                links = links.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
            if fecha_fin:
                links = links.filter(fecha_carga__date__lte=parse_date(fecha_fin))
            if estado:
                if estado.startswith('estado!='):
                    exclude_estado = estado.split('!=')[1]
                    links = links.exclude(estado=exclude_estado)
                else:
                    links = links.filter(estado=estado)
            if categoria_id:
                links = links.filter(categorias__id=categoria_id)
            if tv_digital_id:
                links = links.filter(tv_digital__id=tv_digital_id)
            if solo_propios:
                links = links.filter(cargado_por=request.user)

            for link in links:
                data.append({
                    'id': link.id,
                    'url': link.url,
                    'estado': getattr(link, 'estado', EstadoLink.PENDIENTE),
                    'fecha_carga': link.fecha_carga.isoformat() if link.fecha_carga else None,
                    'revisado_clasificador': getattr(link, 'revisado_clasificador', False),
                    'diario_logo_url': None,
                    'diario_nombre': None,
                    'diario_digital': None,
                    'red_social': None,
                    'red_social_nombre': None,
                    'red_social_logo_url': None,
                    'tv_digital': link.tv_digital.id if link.tv_digital else None,
                    'tv_digital_nombre': link.tv_digital.nombre if link.tv_digital else 'Sin asignar',
                    'tv_digital_logo_url': link.tv_digital.logo.url if link.tv_digital and getattr(link.tv_digital, 'logo', None) else None,
                    'radio_digital': None,
                    'radio_digital_nombre': None,
                    'radio_digital_logo_url': None,
                    'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
                })
        elif fuente == 'radio_digital':
            links = LinkRadioDigital.objects.select_related('radio_digital').prefetch_related('categorias').all()

            if fecha_inicio:
                links = links.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
            if fecha_fin:
                links = links.filter(fecha_carga__date__lte=parse_date(fecha_fin))
            if estado:
                if estado.startswith('estado!='):
                    exclude_estado = estado.split('!=')[1]
                    links = links.exclude(estado=exclude_estado)
                else:
                    links = links.filter(estado=estado)
            if categoria_id:
                links = links.filter(categorias__id=categoria_id)
            if radio_digital_id:
                links = links.filter(radio_digital__id=radio_digital_id)
            if solo_propios:
                links = links.filter(cargado_por=request.user)

            for link in links:
                data.append({
                    'id': link.id,
                    'url': link.url,
                    'estado': getattr(link, 'estado', EstadoLink.PENDIENTE),
                    'fecha_carga': link.fecha_carga.isoformat() if link.fecha_carga else None,
                    'revisado_clasificador': getattr(link, 'revisado_clasificador', False),
                    'diario_logo_url': None,
                    'diario_nombre': None,
                    'diario_digital': None,
                    'red_social': None,
                    'red_social_nombre': None,
                    'red_social_logo_url': None,
                    'tv_digital': None,
                    'tv_digital_nombre': None,
                    'tv_digital_logo_url': None,
                    'radio_digital': link.radio_digital.id if link.radio_digital else None,
                    'radio_digital_nombre': link.radio_digital.nombre if link.radio_digital else 'Sin asignar',
                    'radio_digital_logo_url': link.radio_digital.logo.url if link.radio_digital and getattr(link.radio_digital, 'logo', None) else None,
                    'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
                })
        else:
            # Filtrar links de diarios digitales
            links = LinkRelevante.objects.select_related('diario_digital').prefetch_related('categorias').all()

            if fecha_inicio:
                links = links.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
            if fecha_fin:
                links = links.filter(fecha_carga__date__lte=parse_date(fecha_fin))
            if diario_id:
                links = links.filter(diario_digital__id=diario_id)
            if estado:
                if estado.startswith('estado!='):
                    exclude_estado = estado.split('!=')[1]
                    links = links.exclude(estado=exclude_estado)
                else:
                    links = links.filter(estado=estado)
            if categoria_id:
                links = links.filter(categorias__id=categoria_id)
            if solo_propios:
                links = links.filter(cargado_por=request.user)

            for link in links:
                data.append({
                    'id': link.id,
                    'url': link.url,
                    'estado': getattr(link, 'estado', EstadoLink.PENDIENTE),
                    'fecha_carga': link.fecha_carga.isoformat() if link.fecha_carga else None,
                    'revisado_clasificador': getattr(link, 'revisado_clasificador', False),
                    'diario_logo_url': link.diario_digital.logo.url if hasattr(link.diario_digital, 'logo') and link.diario_digital and link.diario_digital.logo else None,
                    'diario_nombre': link.diario_digital.nombre if link.diario_digital else 'Sin diario',
                    'diario_digital': link.diario_digital.id if link.diario_digital else None,
                    'red_social': None,
                    'red_social_nombre': None,
                    'red_social_logo_url': None,
                    'tv_digital': None,
                    'tv_digital_nombre': None,
                    'tv_digital_logo_url': None,
                    'radio_digital': None,
                    'radio_digital_nombre': None,
                    'radio_digital_logo_url': None,
                    'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
                })

        return JsonResponse(data, safe=False)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_diarios(request):
    try:
        diarios = DiarioDigital.objects.all()
        data = [{'id': d.id, 'nombre': d.nombre} for d in diarios]
        return JsonResponse(data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_categorias(request):
    try:
        categorias = Categoria.objects.all()
        data = [{'id': c.id, 'nombre': c.nombre} for c in categorias]
        return JsonResponse(data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
def api_link_detail(request, link_id):
    try:
        link = LinkRelevante.objects.get(id=link_id)
        
        if request.method == 'PATCH':
            import json
            data = json.loads(request.body)
            
            # Actualizar campos
            if 'estado' in data:
                link.estado = data['estado']
            if 'revisado_clasificador' in data:
                link.revisado_clasificador = data['revisado_clasificador']
            if 'categorias' in data:
                link.categorias.set(data['categorias'])
            
            link.save()
            
            return JsonResponse({
                'id': link.id,
                'estado': link.estado,
                'revisado_clasificador': link.revisado_clasificador,
                'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
            })
        
        # GET request - devolver datos del link
        return JsonResponse({
            'id': link.id,
            'url': link.url,
            'estado': link.estado,
            'revisado_clasificador': link.revisado_clasificador,
            'categorias_info': [{'id': cat.id, 'nombre': cat.nombre} for cat in link.categorias.all()]
        })
        
    except LinkRelevante.DoesNotExist:
        return JsonResponse({'error': 'Link no encontrado'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# -------------------- FUNCTION-BASED API --------------------


@require_GET
@login_required
def api_links(request):
    # La lógica de filtrado de roles está ahora en LinkRelevanteViewSet.get_queryset.
    # Aquí solo aplicamos los filtros y devolvemos los datos.
    estado = request.GET.get("estado")
    fecha_inicio = request.GET.get("fecha_inicio")
    fecha_fin = request.GET.get("fecha_fin")
    categoria_id = request.GET.get("categoria_id")
    diario_id = request.GET.get("diario_id")
    
    qs = LinkRelevante.objects.select_related("diario_digital").prefetch_related(
        "categorias"
    )

    # Lógica de filtro basada en rol (redundante con ViewSet pero mantenida por compatibilidad)
    rol = request.user.userprofile.rol
    if rol == Roles.PRENSA:
        qs = qs.filter(cargado_por=request.user)
    elif rol in [Roles.CLIENTE, Roles.INFORMES]:
        qs = qs.filter(estado='aprobado')

    if estado:
        qs = qs.filter(estado=estado)
    if fecha_inicio:
        qs = qs.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
    if fecha_fin:
        qs = qs.filter(fecha_carga__date__lte=parse_date(fecha_fin))
    if categoria_id:
        qs = qs.filter(categorias__id=categoria_id)
    if diario_id and diario_id.isdigit():
        qs = qs.filter(diario_digital__id=int(diario_id))

    serializer = LinkRelevanteSerializer(qs.distinct().order_by('-fecha_carga'), many=True)
    return JsonResponse(serializer.data, safe=False)


@require_GET
@login_required
def lista_links(request):
    # Esta función parece ser idéntica a api_links pero sin la restricción de rol en el queryset.
    # Se mantendrá la lógica de filtrado de estado y fechas para que sea útil.
    estado = request.GET.get("estado")
    fecha_inicio = request.GET.get("fecha_inicio")
    fecha_fin = request.GET.get("fecha_fin")
    categoria_id = request.GET.get("categoria_id")
    diario_id = request.GET.get("diario_id")
    
    qs = LinkRelevante.objects.select_related("diario_digital").prefetch_related(
        "categorias"
    )

    if estado:
        qs = qs.filter(estado=estado)
    if fecha_inicio:
        qs = qs.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
    if fecha_fin:
        qs = qs.filter(fecha_carga__date__lte=parse_date(fecha_fin))
    if categoria_id:
        qs = qs.filter(categorias__id=categoria_id)
    if diario_id and diario_id.isdigit():
        qs = qs.filter(diario_digital__id=int(diario_id))

    serializer = LinkRelevanteSerializer(qs.distinct().order_by('-fecha_carga'), many=True)
    return JsonResponse(serializer.data, safe=False)


# -------------------- WEB DE INFORMES Y ESTADÍSTICAS --------------------


@login_required
def informes_view(request):
    if request.user.userprofile.rol not in [
        Roles.INFORMES,
        Roles.ADMIN,
        Roles.GERENCIA,
    ]:
        return render(request, "403.html", status=403)

    # ------------------ FILTRO DE ARTÍCULOS ------------------
    articulos = Articulo.objects.select_related("categoria", "generado_por").order_by(
        "-fecha_creacion"
    )

    # Obtener filtros del GET
    fecha_desde = request.GET.get("fecha_desde")
    fecha_hasta = request.GET.get("fecha_hasta")
    categoria_id = request.GET.get("categoria_id")
    generado_por_id = request.GET.get("generado_por_id")

    # Aplicar filtros si existen
    if fecha_desde:
        articulos = articulos.filter(fecha_creacion__date__gte=parse_date(fecha_desde))
    if fecha_hasta:
        articulos = articulos.filter(fecha_creacion__date__lte=parse_date(fecha_hasta))
    if categoria_id:
        articulos = articulos.filter(categoria__id=categoria_id)
    if generado_por_id:
        articulos = articulos.filter(generado_por__id=generado_por_id)

    # ------------------ INFORMES ------------------
    # Solo se cargan los informes que corresponden al usuario, a menos que sea ADMIN o GERENCIA.
    qs_informes = InformeIndividual.objects.select_related("generado_por").order_by(
        "-fecha_creacion"
    )
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA]:
        informes = qs_informes.filter(generado_por=request.user)
    else:
        informes = qs_informes

    # ------------------ PARA LOS SELECTS ------------------
    categorias = Categoria.objects.all()
    # Solo usuarios que pueden generar artículos
    usuarios = User.objects.filter(userprofile__rol__in=[Roles.ADMIN, Roles.REDACCION, Roles.PRENSA]) 

    return render(
        request,
        "informes.html",
        {
            "articulos": articulos,
            "informes": informes,
            "categorias": categorias,
            "usuarios": usuarios,
            "fecha_desde": fecha_desde,
            "fecha_hasta": fecha_hasta,
            "selected_categoria": int(categoria_id) if categoria_id and categoria_id.isdigit() else 0,
            "selected_usuario": int(generado_por_id) if generado_por_id and generado_por_id.isdigit() else 0,
        },
    )


@login_required
def informes_crear_view(request):
    if request.user.userprofile.rol not in [
        Roles.ADMIN,
        Roles.INFORMES,
        Roles.GERENCIA,
    ]:
        return render(request, "403.html", status=403)

    # ------------------ INFORMES (Lista para la tabla lateral) -------------------
    qs_informes = InformeIndividual.objects.select_related("generado_por").order_by(
        "-fecha_creacion"
    )
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA]:
        informes = qs_informes.filter(generado_por=request.user)
    else:
        informes = qs_informes

    # ------------------ ARTÍCULOS CON FILTROS -------------------
    articulos = Articulo.objects.select_related("categoria", "generado_por").order_by(
        "-fecha_creacion"
    )

    fecha_inicio = request.GET.get("fecha_inicio")
    fecha_fin = request.GET.get("fecha_fin")
    categoria_id = request.GET.get("categoria")
    usuario_id = request.GET.get("usuario")

    if fecha_inicio:
        articulos = articulos.filter(fecha_creacion__date__gte=parse_date(fecha_inicio))
    if fecha_fin:
        articulos = articulos.filter(fecha_creacion__date__lte=parse_date(fecha_fin))
    if categoria_id and categoria_id != "0":
        articulos = articulos.filter(categoria__id=categoria_id)
    if usuario_id and usuario_id != "0":
        articulos = articulos.filter(generado_por__id=usuario_id)

    # ------------------ POST NUEVO INFORME -------------------
    if request.method == "POST":
        # Se necesita `request.FILES` para manejar la subida de la foto
        form = InformeIndividualForm(request.POST, request.FILES)
        if form.is_valid():
            informe = form.save(commit=False)
            informe.generado_por = request.user
            
            # Se maneja la foto aquí (el form no lo hace automáticamente con commit=False)
            if "foto" in request.FILES:
                informe.foto = request.FILES["foto"]
                
            informe.save()
            # Asegurarse de que el objeto esté guardado antes de manipular ManyToMany (m2m)

            # 👇 Creación de M2M y modelos relacionados (debe ir DESPUÉS de informe.save())

            # Alias
            for nombre in request.POST.getlist("alias[]"):
                if nombre.strip():
                    informe.alias.add(Alias.objects.create(nombre=nombre))

            # Teléfonos (asumiendo que es M2M o OneToMany - el modelo 'Telefono' lo dicta)
            # Asumiendo que `informe.telefono` es un ManyToManyField a Telefono
            for numero in request.POST.getlist("telefono[]"): 
                if numero.strip():
                    # Aquí asumo que estás intentando crear un objeto Telefono y adjuntarlo
                    # Si fuera OneToMany/ForeignKey, la lógica sería diferente.
                    informe.telefono.add(Telefono.objects.create(numero=numero)) 

            # Domicilios
            for calle, altura, barrio, ciudad, provincia in zip(
                request.POST.getlist("domicilio_calle[]"),
                request.POST.getlist("domicilio_altura[]"),
                request.POST.getlist("domicilio_barrio[]"),
                request.POST.getlist("domicilio_ciudad[]"),
                request.POST.getlist("domicilio_provincia[]"),
            ):
                if any(
                    [
                        calle.strip(),
                        altura.strip(),
                        barrio.strip(),
                        ciudad.strip(),
                        provincia.strip(),
                    ]
                ):
                    informe.domicilio.add(
                        Domicilio.objects.create(
                            calle=calle,
                            altura=altura,
                            barrio=barrio,
                            ciudad=ciudad,
                            provincia=provincia,
                        )
                    )

            # Vehículos
            for marca, modelo, dominio, color in zip(
                request.POST.getlist("vehiculo_marca[]"),
                request.POST.getlist("vehiculo_modelo[]"),
                request.POST.getlist("vehiculo_dominio[]"),
                request.POST.getlist("vehiculo_color[]"),
            ):
                if any([marca.strip(), modelo.strip(), dominio.strip(), color.strip()]):
                    informe.vehiculos.add(
                        Vehiculo.objects.create(
                            marca=marca, modelo=modelo, dominio=dominio, color=color
                        )
                    )

            # Empleadores
            for nombre, cuit, alta, baja in zip(
                request.POST.getlist("empleador_nombre[]"),
                request.POST.getlist("empleador_cuit[]"),
                request.POST.getlist("empleador_fecha_alta[]"),
                request.POST.getlist("empleador_fecha_baja[]"),
            ):
                if any([nombre.strip(), cuit.strip(), alta.strip(), baja.strip()]):
                    informe.empleadores.add(
                        Empleador.objects.create(
                            nombre=nombre,
                            cuit=cuit,
                            fecha_alta=parse_date(alta) if alta else None, # Usar parse_date
                            fecha_baja=parse_date(baja) if baja else None, # Usar parse_date
                        )
                    )

            # Vínculos
            for nombre_apellido, dni, tipo in zip(
                request.POST.getlist("vinculo_nombre_apellido[]"),
                request.POST.getlist("vinculo_dni[]"),
                request.POST.getlist("vinculo_tipo[]"),
            ):
                if any([nombre_apellido.strip(), dni.strip(), tipo.strip()]):
                    informe.vinculos.add(
                        Vinculo.objects.create(
                            nombre_apellido=nombre_apellido, dni=dni, tipo=tipo
                        )
                    )

            # Artículos
            for art_id in request.POST.getlist("articulos[]"):
                try:
                    informe.articulos.add(Articulo.objects.get(id=art_id))
                except (Articulo.DoesNotExist, ValueError):
                    continue

            log_actividad(
                request,
                TipoActividad.CARGA_INFORME,
                f"Se creó un informe individual ID {informe.id}",
            )
            messages.success(request, f"Informe ID {informe.id} creado correctamente.")
            return redirect("crear_informe")
        else:
            logger.error("Errores del formulario de InformeIndividual:", form.errors)
    else:
        form = InformeIndividualForm()

    # ------------------ CONTEXT -------------------
    return render(
        request,
        "informes.html",
        {
            "form": form,
            "informes": informes,
            "articulos": articulos,
            "categorias": Categoria.objects.all(),
            "usuarios": User.objects.all(),
            "selected_categoria": (
                int(categoria_id) if categoria_id and categoria_id.isdigit() and categoria_id != "0" else 0
            ),
            "selected_usuario": (
                int(usuario_id) if usuario_id and usuario_id.isdigit() and usuario_id != "0" else 0
            ),
            "fecha_inicio": fecha_inicio,
            "fecha_fin": fecha_fin,
        },
    )

# -------------------- INFORMES INDIVIDUALES (Parte 2) --------------------


@login_required
def api_detalle_informe(request, id):
    # CORRECCIÓN: Se usa get_object_or_404 y prefetch de una vez.
    try:
        informe = get_object_or_404(
            InformeIndividual.objects.select_related("generado_por").prefetch_related(
                "alias",
                "telefono",
                "domicilio",
                "vehiculos",
                "empleadores",
                "vinculos",
                "articulos__categoria", # Añadir la categoría de los artículos
            ),
            id=id
        )
    except Exception as e:
        logger.error(f"Error al obtener detalle de informe ID {id}: {e}")
        return JsonResponse({"status": "error", "msg": "Error interno al procesar el informe."}, status=500)

    # Armamos manualmente el diccionario
    data = {
        "id": informe.id,
        "apellido": informe.apellido,
        "nombre": informe.nombre,
        "documento": informe.documento,
        "cuit": informe.cuit,
        "nacionalidad": informe.nacionalidad,
        "banda": informe.banda,
        "rol": informe.get_rol_display(),
        "situacion": informe.get_situacion_display(),
        "actividad": informe.actividad,
        "foto": informe.foto.url if informe.foto else None,
        "fecha_nacimiento": (
            informe.fecha_nacimiento.strftime("%d/%m/%Y")
            if informe.fecha_nacimiento
            else ""
        ),
        "generado_por": informe.generado_por.username if informe.generado_por else "-",
        "fecha_creacion": informe.fecha_creacion.strftime("%d/%m/%Y %H:%M"),
        "fecha_modificacion": informe.fecha_modificacion.strftime("%d/%m/%Y %H:%M"),
        "alias": [a.nombre for a in informe.alias.all()],
        "telefonos": [t.numero for t in informe.telefono.all()],
        "domicilios": [
            {
                "calle": d.calle,
                "altura": d.altura,
                "barrio": d.barrio,
                "ciudad": d.ciudad,
                "provincia": d.provincia,
            }
            for d in informe.domicilio.all()
        ],
        "vehiculos": [
            {
                "marca": v.marca,
                "modelo": v.modelo,
                "dominio": v.dominio,
                "color": v.color,
            }
            for v in informe.vehiculos.all()
        ],
        "empleadores": [
            {
                "nombre": e.nombre,
                "cuit": e.cuit,
                # CORRECCIÓN: Manejar fechas nulas
                "alta": e.fecha_alta.strftime("%d/%m/%Y") if e.fecha_alta else "", 
                "baja": e.fecha_baja.strftime("%d/%m/%Y") if e.fecha_baja else "",
            }
            for e in informe.empleadores.all()
        ],
        "vinculos": [
            {"nombre_apellido": v.nombre_apellido, "dni": v.dni, "tipo": v.tipo}
            for v in informe.vinculos.all()
        ],
        "articulos": [
            {
                "titulo": a.titulo,
                "categoria": a.categoria.nombre if a.categoria else "-",
                "fecha": a.fecha_creacion.strftime("%d/%m/%Y %H:%M"),
                "id": a.id, # Añadir el ID para referencia
            }
            for a in informe.articulos.all()
        ],
    }
    return JsonResponse({"status": "ok", "informe": data})


@login_required
def eliminar_informe_view(request, id):
    # Se añade restricción de rol (ADMIN/INFORMES/GERENCIA) para eliminar
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.INFORMES, Roles.GERENCIA]:
        return render(request, "403.html", status=403)
        
    informe = get_object_or_404(InformeIndividual, id=id)
    if request.method == "POST":
        informe.delete()
        log_actividad(request, TipoActividad.OTRO, f"Se eliminó el informe ID {id}")
        messages.success(request, f"Informe ID {id} eliminado correctamente.")
        # Redirigir a la vista de lista principal de informes
        return redirect("informes") 
    
    # Si es GET, muestra la página de confirmación
    return render(request, "confirm_delete.html", {"informe": informe, "tipo": "Informe Individual"})


@login_required
def consulta_informes_view(request):
    # Esta vista debe ser accesible para CLIENTE, ADMIN, GERENCIA
    if request.user.userprofile.rol not in [Roles.CLIENTE, Roles.ADMIN, Roles.GERENCIA]:
        return render(request, "403.html", status=403)
        
    informes = InformeIndividual.objects.select_related(
        "generado_por"
    ).prefetch_related(
        "alias",
        "telefono",
        "domicilio",
        "vehiculos",
        "empleadores",
        "vinculos",
        "articulos",
    )

    # Recoger los filtros del GET
    apellido = request.GET.get("apellido", "").strip()
    nombre = request.GET.get("nombre", "").strip()
    documento = request.GET.get("documento", "").strip()
    nacionalidad = request.GET.get("nacionalidad", "").strip()
    banda = request.GET.get("banda", "").strip()
    actividad = request.GET.get("actividad", "").strip()
    rol = request.GET.get("rol", "").strip()
    situacion = request.GET.get("situacion", "").strip()
    alias = request.GET.get("alias", "").strip()
    telefono = request.GET.get("telefono", "").strip()
    calle = request.GET.get("calle", "").strip()
    ciudad = request.GET.get("ciudad", "").strip()
    provincia = request.GET.get("provincia", "").strip()
    vehiculo_marca = request.GET.get("vehiculo_marca", "").strip()
    vehiculo_dominio = request.GET.get("vehiculo_dominio", "").strip()
    vehiculo_color = request.GET.get("vehiculo_color", "").strip()
    empleador_nombre = request.GET.get("empleador_nombre", "").strip()
    empleador_cuit = request.GET.get("empleador_cuit", "").strip()
    vinculo_nombre = request.GET.get("vinculo_nombre", "").strip()
    vinculo_dni = request.GET.get("vinculo_dni", "").strip()
    altura_desde = request.GET.get("altura_desde")
    altura_hasta = request.GET.get("altura_hasta")

    fecha_nacimiento_desde = request.GET.get("fecha_nacimiento_desde", "").strip()
    fecha_nacimiento_hasta = request.GET.get("fecha_nacimiento_hasta", "").strip()

    # Filtros de altura de domicilio
    # CORRECCIÓN: Se debe asegurar que el campo `altura` de Domicilio es un campo numérico
    # Si es un CharField, el filtro no funcionará con `__gte` y `__lte`
    # Asumiendo que es un campo numérico:
    if altura_desde and altura_desde.isdigit():
        informes = informes.filter(domicilio__altura__gte=int(altura_desde))
    if altura_hasta and altura_hasta.isdigit():
        informes = informes.filter(domicilio__altura__lte=int(altura_hasta))

    # Filtros directos
    if apellido:
        informes = informes.filter(apellido__icontains=apellido)
    if nombre:
        informes = informes.filter(nombre__icontains=nombre)
    if documento:
        informes = informes.filter(documento__icontains=documento)
    if nacionalidad:
        informes = informes.filter(nacionalidad__icontains=nacionalidad)
    if banda:
        informes = informes.filter(banda__icontains=banda)
    if actividad:
        informes = informes.filter(actividad__icontains=actividad)
    if rol:
        informes = informes.filter(rol=rol)
    if situacion:
        informes = informes.filter(situacion=situacion)

    # Filtros indirectos por relaciones
    if alias:
        informes = informes.filter(alias__nombre__icontains=alias)
    if telefono:
        informes = informes.filter(telefono__numero__icontains=telefono)
    if calle:
        informes = informes.filter(domicilio__calle__icontains=calle)
    if ciudad:
        informes = informes.filter(domicilio__ciudad__icontains=ciudad)
    if provincia:
        informes = informes.filter(domicilio__provincia__icontains=provincia)
    if vehiculo_marca:
        informes = informes.filter(vehiculos__marca__icontains=vehiculo_marca)
    if vehiculo_dominio:
        informes = informes.filter(vehiculos__dominio__icontains=vehiculo_dominio)
    if vehiculo_color:
        informes = informes.filter(vehiculos__color__icontains=vehiculo_color)
    if empleador_nombre:
        informes = informes.filter(empleadores__nombre__icontains=empleador_nombre)
    if empleador_cuit:
        informes = informes.filter(empleadores__cuit__icontains=empleador_cuit)
    if vinculo_nombre:
        informes = informes.filter(vinculos__nombre_apellido__icontains=vinculo_nombre)
    if vinculo_dni:
        informes = informes.filter(vinculos__dni__icontains=vinculo_dni)

    # Fecha nacimiento aproximada
    if fecha_nacimiento_desde:
        fecha_desde = parse_date(fecha_nacimiento_desde)
        if fecha_desde:
            informes = informes.filter(fecha_nacimiento__gte=fecha_desde)

    if fecha_nacimiento_hasta:
        fecha_hasta = parse_date(fecha_nacimiento_hasta)
        if fecha_hasta:
            informes = informes.filter(fecha_nacimiento__lte=fecha_hasta)

    informes = informes.distinct().order_by("-fecha_creacion")

    return render(request, "consulta_informes.html", {"informes": informes, "Roles": Roles})


@login_required
def editar_individuo_view(request, id):
    # CORRECCIÓN: La vista original estaba mezclando crear y editar. La forma más limpia
    # es usar esta función solo para editar. La creación se hace en informes_crear_view (POST).
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.INFORMES, Roles.GERENCIA]:
        return render(request, "403.html", status=403)
        
    # El `id` es obligatorio para editar
    individuo = get_object_or_404(InformeIndividual, pk=id)
    
    if request.method == "POST":
        # Se necesita request.FILES para la foto
        form = InformeIndividualForm(request.POST, request.FILES, instance=individuo)
        if form.is_valid():
            # TODO: Aquí se necesita manejar la actualización de los modelos relacionados (Alias, Domicilios, etc)
            # de forma similar a `informes_crear_view`. 
            # Si sólo se usa `form.save()`, los modelos relacionados NO se actualizarán.
            form.save() 
            messages.success(request, f"Informe ID {individuo.id} actualizado correctamente.")
            return redirect("informes")
        else:
            logger.error("Errores al editar informe:", form.errors)
    else:
        form = InformeIndividualForm(instance=individuo)
        
    return render(
        request, "editar_individuo.html", {"form": form, "individuo": individuo}
    )


@login_required
def estadisticas_view(request):
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA]:
        return render(request, "403.html", status=403)

    usuarios = User.objects.filter(
        userprofile__rol__in=[Roles.ADMIN, Roles.PRENSA, Roles.CLASIFICACION]
    )
    links = LinkRelevante.objects.select_related("cargado_por").prefetch_related(
        "categorias"
    )

    usuario = request.GET.get("usuario")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")

    if usuario:
        links = links.filter(cargado_por__username=usuario)
    if desde:
        links = links.filter(fecha_carga__date__gte=parse_date(desde))
    if hasta:
        links = links.filter(fecha_carga__date__lte=parse_date(hasta))

    # ----------------------------
    # PROMEDIO ENTRE CARGA Y APROBACIÓN
    # ----------------------------
    # Aseguramos que solo se calculen para links con fecha_aprobacion
    links_aprobados = links.filter(estado='aprobado').exclude(fecha_aprobacion__isnull=True).annotate(
        tiempo_aprobacion=ExpressionWrapper(
            F("fecha_aprobacion") - F("fecha_carga"), output_field=DurationField()
        )
    )
    promedio = links_aprobados.aggregate(promedio=Avg("tiempo_aprobacion"))["promedio"]

    promedio_dias = round(promedio.total_seconds() / 86400, 2) if promedio else 0
    promedio_horas = round(promedio.total_seconds() / 3600, 2) if promedio else 0

    # ----------------------------
    # GRÁFICOS
    # ----------------------------
    # Conteo por estado
    estado_count = links.values("estado").annotate(cantidad=Count("id"))
    torta = {item["estado"]: item["cantidad"] for item in estado_count}

    # Conteo por día de carga
    barras = defaultdict(int)
    for link in links:
        dia = link.fecha_carga.date().isoformat()
        barras[dia] += 1
        
    # Conteo por categoría
    categorias_counter = defaultdict(int)
    for link in links:
        for cat in link.categorias.all():
            categorias_counter[cat.nombre] += 1


    return render(
        request,
        "estadisticas.html",
        {
            "usuarios": usuarios,
            "tabla": links.order_by("-fecha_carga"),
            "torta": torta,
            "barras_labels": list(barras.keys()),
            "barras_values": list(barras.values()),
            "categorias_labels": list(categorias_counter.keys()),
            "categorias_values": list(categorias_counter.values()),
            "filtros": {"usuario": usuario, "desde": desde, "hasta": hasta},
            "promedio_dias": promedio_dias,
            "promedio_horas": promedio_horas,
        },
    )


@login_required
def exportar_estadisticas_pdf(request):
    if HTML is None:
        return _weasyprint_unavailable_response(request, "estadisticas")
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA]:
        return render(request, "403.html", status=403)
        
    usuario = request.GET.get("usuario")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")

    links = LinkRelevante.objects.select_related("cargado_por").prefetch_related(
        "categorias"
    )
    if usuario:
        links = links.filter(cargado_por__username=usuario)
    if desde:
        links = links.filter(fecha_carga__date__gte=parse_date(desde))
    if hasta:
        links = links.filter(fecha_carga__date__lte=parse_date(hasta))

    # ----------------------------
    # Gráficos Plotly
    # ----------------------------

    # Gráfico de barras (Links por Día)
    data_barras = (
        links.extra(select={"day": "date(fecha_carga)"})
        .values("day")
        .order_by("day")
        .annotate(count=Count("id"))
    )
    df_barras = pd.DataFrame(data_barras)
    fig_barras = px.bar(
        df_barras,
        x="day",
        y="count",
        labels={"day": "Fecha", "count": "Cantidad"},
        title="Links por Día",
    )
    buf_barras = io.BytesIO()
    # ⚠️ REQUiere KALEIDO instalado para exportar a PNG ⚠️
    try:
        fig_barras.write_image(buf_barras, format="png")
    except ValueError as e:
        logger.error(f"Error al generar gráfico de barras (Plotly/Kaleido): {e}")
        # Manejo de error para evitar 500
        barra_img = None 
    else:
        buf_barras.seek(0)
        barra_img = base64.b64encode(buf_barras.read()).decode("utf-8")


    # Gráfico de torta (Distribución por Estado)
    estado_data = links.values("estado").annotate(count=Count("id"))
    df_torta = pd.DataFrame(estado_data)
    fig_torta = px.pie(
        df_torta, names="estado", values="count", title="Distribución por Estado"
    )
    buf_torta = io.BytesIO()
    try:
        fig_torta.write_image(buf_torta, format="png")
    except ValueError as e:
        logger.error(f"Error al generar gráfico de torta (Plotly/Kaleido): {e}")
        torta_img = None 
    else:
        buf_torta.seek(0)
        torta_img = base64.b64encode(buf_torta.read()).decode("utf-8")


    # Gráfico de categorías
    categoria_counter = defaultdict(int)
    for link in links:
        for cat in link.categorias.all():
            categoria_counter[cat.nombre] += 1
            
    df_cats = pd.DataFrame(
        {
            "categoria": list(categoria_counter.keys()),
            "cantidad": list(categoria_counter.values()),
        }
    )
    fig_cats = px.bar(df_cats, x="categoria", y="cantidad", title="Links por Categoría")
    buf_cats = io.BytesIO()
    try:
        fig_cats.write_image(buf_cats, format="png")
    except ValueError as e:
        logger.error(f"Error al generar gráfico de categorías (Plotly/Kaleido): {e}")
        cats_img = None
    else:
        buf_cats.seek(0)
        cats_img = base64.b64encode(buf_cats.read()).decode("utf-8")
        
    # ----------------------------
    # Logo en PDF
    # ----------------------------
    logo_path = os.path.join(settings.BASE_DIR, "static", "img", "logo.png")
    logo_b64 = None
    try:
        with open(logo_path, "rb") as img_f:
            logo_b64 = base64.b64encode(img_f.read()).decode("utf-8")
    except FileNotFoundError:
        logger.error(f"No se encontró el archivo de logo en: {logo_path}")
        # Si no se encuentra el logo, se pasa None y el template debe manejarlo
    except Exception as e:
        logger.error(f"Error al codificar el logo en base64: {e}")
        
        
    # ----------------------------
    # Renderizado y Exportación
    # ----------------------------

    html = render_to_string(
        "estadisticas_pdf.html",
        {
            "tabla": links.order_by("-fecha_carga"),
            "barra_img": barra_img,
            "torta_img": torta_img,
            "cats_img": cats_img,
            "usuario": usuario,
            "desde": desde,
            "hasta": hasta,
            "generado": datetime.now(),
            "logo_b64": logo_b64,
        },
    )
    
    try:
        # Asegurarse que WeasyPrint tenga una URL base para resolver recursos
        pdf = HTML(string=html, base_url=request.build_absolute_uri()).write_pdf()
    except Exception as e:
        logger.error(f"Error al generar el PDF con WeasyPrint: {e}")
        messages.error(request, f"Error al generar PDF. Verifique dependencias (WeasyPrint, Kaleido). Detalle: {e}")
        return redirect('estadisticas') # Redirigir de vuelta para evitar el 500

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="estadisticas_links.pdf"'
    return response


@login_required
@require_GET
def estadisticas_api_view(request):
    # EndPoint JSON para estadísticas
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA]:
        return JsonResponse({"error": "Acceso denegado"}, status=403)
        
    usuario = request.GET.get("usuario")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")
    
    links = LinkRelevante.objects.select_related("cargado_por").prefetch_related(
        "categorias"
    )
    
    if usuario:
        links = links.filter(cargado_por__username=usuario)
    if desde:
        links = links.filter(fecha_carga__date__gte=parse_date(desde))
    if hasta:
        links = links.filter(fecha_carga__date__lte=parse_date(hasta))
        
    # Conteo por estado
    estado_count = links.values("estado").annotate(cantidad=Count("id"))
    torta = {x["estado"]: x["cantidad"] for x in estado_count}
    
    # Conteo por día
    barras = defaultdict(int)
    for link in links:
        barras[link.fecha_carga.date().isoformat()] += 1
        
    # Conteo por categoría
    categorias = defaultdict(int)
    for link in links:
        for cat in link.categorias.all():
            categorias[cat.nombre] += 1
            
    # Datos de la tabla (limitados)
    tabla = [
        {
            "fecha": l.fecha_carga.strftime("%Y-%m-%d %H:%M"),
            "url": l.url,
            "estado": l.get_estado_display(),
            "usuario": l.cargado_por.username,
        }
        for l in links.order_by('-fecha_carga')[:100]
    ]
    
    # Cálculo de promedio de aprobación
    links_aprobados = links.filter(estado='aprobado').exclude(fecha_aprobacion__isnull=True).annotate(
        tiempo_aprobacion=ExpressionWrapper(
            F("fecha_aprobacion") - F("fecha_carga"), output_field=DurationField()
        )
    )
    promedio = links_aprobados.aggregate(promedio=Avg("tiempo_aprobacion"))["promedio"]
    promedio_horas = round(promedio.total_seconds() / 3600, 2) if promedio else 0

    return JsonResponse(
        {
            "torta": torta,
            "barras_labels": list(barras.keys()),
            "barras_values": list(barras.values()),
            "categorias_labels": list(categorias.keys()),
            "categorias_values": list(categorias.values()),
            "tabla": tabla,
            "promedio_horas": promedio_horas
        }
    )


@login_required
def consultar_articulo_view(request, id):
    # Se debe permitir a REDACCION, ADMIN y CLIENTE consultar artículos.
    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.REDACCION, Roles.CLIENTE]:
        return JsonResponse({"status": "error", "message": "Acceso denegado"}, status=403)
        
    try:
        articulo = Articulo.objects.select_related('categoria').get(pk=id)
        
        # Marcar como consultado (asumiendo que 'consultado' es un campo para CLIENTE)
        if request.user.userprofile.rol == Roles.CLIENTE:
            articulo.consultado = True
            articulo.save(update_fields=['consultado'])
        
        # log_actividad(request, TipoActividad.OTRO, f"Consulta de artículo ID {id}")
        
        data = {
            "titulo": articulo.titulo,
            "descripcion": articulo.descripcion,
            "categoria": articulo.categoria.nombre if articulo.categoria else "-",
            "fecha": articulo.fecha_creacion.strftime("%d/%m/%Y %H:%M"),
        }
        return JsonResponse({"status": "ok", "articulo": data})
    except Articulo.DoesNotExist:
        return JsonResponse(
            {"status": "error", "message": "Artículo no encontrado"}, status=404
        )
    except Exception as e:
        logger.error(f"Error al consultar artículo ID {id}: {e}")
        return JsonResponse({"status": "error", "message": "Error interno"}, status=500)


@login_required
def configuraciones(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        # CORRECCIÓN: Usar la plantilla 403.html
        return render(request, "403.html", status=403)
        
    categorias = Categoria.objects.all()
    diarios = DiarioDigital.objects.all()
    redes_sociales = RedSocial.objects.all()
    tv_digital = TvDigital.objects.all()
    radios_digitales = RadioDigital.objects.all()
    return render(
        request,
        "configuraciones.html",
        {
            "categorias": categorias,
            "diarios": diarios,
            "redes_sociales": redes_sociales,
            "tv_digital": tv_digital,
            "radios_digitales": radios_digitales,
        },
    )


# ========================
# CRUD Herramientas OSINT
# ========================


@login_required
def herramientas_osint_view(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    herramientas = HerramientaOSINT.objects.all()
    return render(
        request,
        "herramientas_osint.html",
        {"herramientas": herramientas},
    )


@login_required
def crear_herramienta_osint(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = HerramientaOSINTForm(request.POST, request.FILES)
        if form.is_valid():
            herramienta = form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Herramienta OSINT creada: {herramienta.nombre}",
            )
            messages.success(request, "Herramienta guardada correctamente.")
            return redirect("herramientas_osint")
    else:
        form = HerramientaOSINTForm()

    return render(
        request,
        "form_herramienta_osint.html",
        {"form": form, "action": "Crear"},
    )


@login_required
def editar_herramienta_osint(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    herramienta = get_object_or_404(HerramientaOSINT, id=id)

    if request.method == "POST":
        form = HerramientaOSINTForm(request.POST, request.FILES, instance=herramienta)
        if form.is_valid():
            form.save()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Herramienta OSINT editada: {herramienta.nombre}",
            )
            messages.success(request, "Herramienta actualizada correctamente.")
            return redirect("herramientas_osint")
    else:
        form = HerramientaOSINTForm(instance=herramienta)

    return render(
        request,
        "form_herramienta_osint.html",
        {"form": form, "action": "Editar", "herramienta": herramienta},
    )


@login_required
def eliminar_herramienta_osint(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    herramienta = get_object_or_404(HerramientaOSINT, id=id)
    log_actividad(
        request,
        TipoActividad.OTRO,
        f"Herramienta OSINT eliminada: {herramienta.nombre}",
    )
    herramienta.delete()
    messages.success(request, "Herramienta eliminada correctamente.")
    return redirect("herramientas_osint")


@login_required
def exportar_fuentes_osint(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Herramientas OSINT"
    ws.append(["Nombre", "URL", "Tipo", "Descripción"])

    for herramienta in HerramientaOSINT.objects.all().order_by("nombre"):
        ws.append(
            [
                herramienta.nombre,
                herramienta.url,
                herramienta.get_tipo_display(),
                herramienta.descripcion or "",
            ]
        )

    def _agregar_fuentes(titulo, queryset):
        hoja = wb.create_sheet(titulo)
        hoja.append(["Nombre", "URL principal"])
        for fuente in queryset:
            hoja.append([fuente.nombre, fuente.url_principal])

    _agregar_fuentes("Diarios digitales", DiarioDigital.objects.all().order_by("nombre"))
    _agregar_fuentes("Redes sociales", RedSocial.objects.all().order_by("nombre"))
    _agregar_fuentes("TV digital", TvDigital.objects.all().order_by("nombre"))
    _agregar_fuentes("Radios digitales", RadioDigital.objects.all().order_by("nombre"))

    ahora = localtime(now()).strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"fuentes_osint_{ahora}.xlsx"

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{nombre_archivo}"'
    wb.save(response)
    return response


# ========================
# Exportar Links Relevantes
# ========================


@login_required
def exportar_links_relevantes(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    fecha_desde = request.GET.get("desde")
    fecha_hasta = request.GET.get("hasta")
    diario_id = request.GET.get("diario")
    categoria_id = request.GET.get("categoria")

    links = (
        LinkRelevante.objects.filter(
            estado=EstadoLink.APROBADO, revisado_clasificador=True
        )
        .select_related("diario_digital", "cargado_por")
        .prefetch_related("categorias")
        .order_by("-fecha_carga")
    )

    if fecha_desde:
        links = links.filter(fecha_carga__date__gte=parse_date(fecha_desde))
    if fecha_hasta:
        links = links.filter(fecha_carga__date__lte=parse_date(fecha_hasta))
    if diario_id:
        links = links.filter(diario_digital_id=diario_id)
    if categoria_id:
        links = links.filter(categorias__id=categoria_id)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Links Relevantes"
    ws.append(
        [
            "Fecha de carga",
            "URL",
            "Estado",
            "Diario digital",
            "Categorías",
            "Cargado por",
        ]
    )

    for link in links:
        categorias = ", ".join(cat.nombre for cat in link.categorias.all()) or "-"
        ws.append(
            [
                link.fecha_carga.strftime("%Y-%m-%d %H:%M") if link.fecha_carga else "-",
                link.url,
                link.get_estado_display(),
                link.diario_digital.nombre if link.diario_digital else "Sin diario",
                categorias,
                link.cargado_por.username if link.cargado_por else "Desconocido",
            ]
        )

    nombre_archivo = "links_relevantes.xlsx"
    if fecha_desde or fecha_hasta:
        nombre_archivo = f"links_{fecha_desde or 'inicio'}_{fecha_hasta or 'hoy'}.xlsx"
    elif diario_id:
        nombre_archivo = f"links_diario_{diario_id}.xlsx"
    elif categoria_id:
        nombre_archivo = f"links_categoria_{categoria_id}.xlsx"

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{nombre_archivo}"'
    wb.save(response)
    return response


# ========================
# CRUD Categorías
# ========================


@login_required
def crear_categoria(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    if request.method == "POST":
        form = CategoriaForm(request.POST)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Categoría creada: {form.cleaned_data['nombre']}")
            messages.success(request, "Categoría creada exitosamente.")
            return redirect("configuraciones")
    else:
        form = CategoriaForm()
        
    return render(request, "form_categoria.html", {"form": form, "action": "Crear"})


@login_required
def editar_categoria(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    categoria = get_object_or_404(Categoria, id=id)
    if request.method == "POST":
        form = CategoriaForm(request.POST, instance=categoria)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Categoría editada: {categoria.nombre}")
            messages.success(request, "Categoría actualizada exitosamente.")
            return redirect("configuraciones")
    else:
        form = CategoriaForm(instance=categoria)
        
    return render(request, "form_categoria.html", {"form": form, "action": "Editar"})


@login_required
def eliminar_categoria(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    categoria = get_object_or_404(Categoria, id=id)
    log_actividad(request, TipoActividad.OTRO, f"Categoría eliminada: {categoria.nombre}")
    categoria.delete()
    messages.success(request, "Categoría eliminada exitosamente.")
    return redirect("configuraciones")


# ========================
# CRUD Diarios
# ========================


@login_required
def crear_diario(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    if request.method == "POST":
        form = DiarioForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Diario creado: {form.cleaned_data['nombre']}")
            messages.success(request, "Diario creado exitosamente.")
            return redirect("configuraciones")
    else:
        form = DiarioForm()
        
    return render(request, "form_diario.html", {"form": form, "action": "Crear"})


@login_required
def editar_diario(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    diario = get_object_or_404(DiarioDigital, id=id)
    if request.method == "POST":
        form = DiarioForm(request.POST, request.FILES, instance=diario)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Diario editado: {diario.nombre}")
            messages.success(request, "Diario actualizado exitosamente.")
            return redirect("configuraciones")
    else:
        form = DiarioForm(instance=diario)
        
    return render(request, "form_diario.html", {"form": form, "action": "Editar"})


@login_required
def eliminar_diario(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)
        
    diario = get_object_or_404(DiarioDigital, id=id)
    log_actividad(request, TipoActividad.OTRO, f"Diario eliminado: {diario.nombre}")
    diario.delete()
    messages.success(request, "Diario eliminado exitosamente.")
    return redirect("configuraciones")


@login_required
def crear_red_social(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = RedSocialForm(request.POST, request.FILES)
        if form.is_valid():
            red = form.save()
            log_actividad(request, TipoActividad.OTRO, f"Red social creada: {red.nombre}")
            messages.success(request, "Red social creada exitosamente.")
            return redirect("configuraciones")
    else:
        form = RedSocialForm()

    return render(request, "form_diario.html", {"form": form, "action": "Crear", "titulo": "Red Social"})


@login_required
def editar_red_social(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    red = get_object_or_404(RedSocial, id=id)
    if request.method == "POST":
        form = RedSocialForm(request.POST, request.FILES, instance=red)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Red social editada: {red.nombre}")
            messages.success(request, "Red social actualizada exitosamente.")
            return redirect("configuraciones")
    else:
        form = RedSocialForm(instance=red)

    return render(request, "form_diario.html", {"form": form, "action": "Editar", "titulo": "Red Social"})


@login_required
def eliminar_red_social(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    red = get_object_or_404(RedSocial, id=id)
    log_actividad(request, TipoActividad.OTRO, f"Red social eliminada: {red.nombre}")
    red.delete()
    messages.success(request, "Red social eliminada exitosamente.")
    return redirect("configuraciones")


@login_required
def crear_tv_digital(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = TvDigitalForm(request.POST, request.FILES)
        if form.is_valid():
            canal = form.save()
            log_actividad(request, TipoActividad.OTRO, f"TV digital creada: {canal.nombre}")
            messages.success(request, "Canal de TV digital creado exitosamente.")
            return redirect("configuraciones")
    else:
        form = TvDigitalForm()

    return render(request, "form_diario.html", {"form": form, "action": "Crear", "titulo": "TV Digital"})


@login_required
def editar_tv_digital(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    canal = get_object_or_404(TvDigital, id=id)
    if request.method == "POST":
        form = TvDigitalForm(request.POST, request.FILES, instance=canal)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"TV digital editada: {canal.nombre}")
            messages.success(request, "Canal de TV digital actualizado exitosamente.")
            return redirect("configuraciones")
    else:
        form = TvDigitalForm(instance=canal)

    return render(request, "form_diario.html", {"form": form, "action": "Editar", "titulo": "TV Digital"})


@login_required
def eliminar_tv_digital(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    canal = get_object_or_404(TvDigital, id=id)
    log_actividad(request, TipoActividad.OTRO, f"TV digital eliminada: {canal.nombre}")
    canal.delete()
    messages.success(request, "Canal de TV digital eliminado exitosamente.")
    return redirect("configuraciones")


@login_required
def crear_radio_digital(request):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = RadioDigitalForm(request.POST, request.FILES)
        if form.is_valid():
            radio = form.save()
            log_actividad(request, TipoActividad.OTRO, f"Radio digital creada: {radio.nombre}")
            messages.success(request, "Radio digital creada exitosamente.")
            return redirect("configuraciones")
    else:
        form = RadioDigitalForm()

    return render(request, "form_diario.html", {"form": form, "action": "Crear", "titulo": "Radio Digital"})


@login_required
def editar_radio_digital(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    radio = get_object_or_404(RadioDigital, id=id)
    if request.method == "POST":
        form = RadioDigitalForm(request.POST, request.FILES, instance=radio)
        if form.is_valid():
            form.save()
            log_actividad(request, TipoActividad.OTRO, f"Radio digital editada: {radio.nombre}")
            messages.success(request, "Radio digital actualizada exitosamente.")
            return redirect("configuraciones")
    else:
        form = RadioDigitalForm(instance=radio)

    return render(request, "form_diario.html", {"form": form, "action": "Editar", "titulo": "Radio Digital"})


@login_required
def eliminar_radio_digital(request, id):
    if request.user.userprofile.rol != Roles.GERENTE_PRODUCCION:
        return render(request, "403.html", status=403)

    radio = get_object_or_404(RadioDigital, id=id)
    log_actividad(request, TipoActividad.OTRO, f"Radio digital eliminada: {radio.nombre}")
    radio.delete()
    messages.success(request, "Radio digital eliminada exitosamente.")
    return redirect("configuraciones")


@login_required
def crear_hecho_delictivo(request):
    if request.method != "POST":
        return redirect("hechos_delictivos")

    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        messages.error(request, "No tenés permisos para registrar hechos delictivos.")
        return redirect("hechos_delictivos")

    form = HechoDelictivoForm(request.POST, user=request.user)
    if form.is_valid():
        hecho = form.save(commit=False)
        hecho.creado_por = request.user
        hecho.save()
        form.save_m2m()
        request.session.pop("hecho_delictivo_form_data", None)
        log_actividad(
            request,
            TipoActividad.OTRO,
            f"Hecho delictivo creado (ID {hecho.pk}).",
        )
        messages.success(request, "Hecho delictivo registrado correctamente.")
    else:
        request.session["hecho_delictivo_form_data"] = request.POST.urlencode()
        errores = "; ".join(
            [f"{campo}: {', '.join(map(str, lista))}" for campo, lista in form.errors.items()]
        )
        messages.error(request, f"No se pudo registrar el hecho delictivo: {errores}")

    return redirect("hechos_delictivos")


@login_required
def editar_hecho_delictivo(request, id):
    hecho = get_object_or_404(HechoDelictivo, id=id)
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)
    if hecho.creado_por not in [request.user, None] and request.user.userprofile.rol != Roles.ADMIN:
        return render(request, "403.html", status=403)

    if request.method == "POST":
        form = HechoDelictivoForm(request.POST, instance=hecho, user=request.user)
        if form.is_valid():
            hecho = form.save(commit=False)
            if hecho.creado_por is None:
                hecho.creado_por = request.user
            hecho.save()
            form.save_m2m()
            log_actividad(
                request,
                TipoActividad.OTRO,
                f"Hecho delictivo editado (ID {hecho.pk}).",
            )
            messages.success(request, "Hecho delictivo actualizado correctamente.")
            return redirect("hechos_delictivos")
    else:
        form = HechoDelictivoForm(instance=hecho, user=request.user)

    return render(
        request,
        "form_hecho_delictivo.html",
        {
            "form": form,
            "action": "Editar",
            "hecho": hecho,
        },
    )


@login_required
def eliminar_hecho_delictivo(request, id):
    hecho = get_object_or_404(HechoDelictivo, id=id)
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)
    if hecho.creado_por not in [request.user, None] and request.user.userprofile.rol != Roles.ADMIN:
        return render(request, "403.html", status=403)

    log_actividad(
        request,
        TipoActividad.OTRO,
        f"Hecho delictivo eliminado (ID {hecho.pk}).",
    )
    hecho.delete()
    messages.success(request, "Hecho delictivo eliminado correctamente.")
    return redirect("hechos_delictivos")


@login_required
def hecho_delictivo_detalle_view(request, id):
    hecho = get_object_or_404(
        HechoDelictivo.objects.select_related("creado_por", "articulo").prefetch_related(
            "autor", "noticias"
        ),
        id=id,
    )
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)
    if (
        request.user.userprofile.rol == Roles.REDACCION
        and hecho.creado_por not in [request.user, None]
    ):
        return render(request, "403.html", status=403)

    return render(
        request,
        "hecho_delictivo_detalle.html",
        {
            "hecho": hecho,
        },
    )


@login_required
def exportar_hecho_delictivo_pdf(request, id):
    if HTML is None:
        return _weasyprint_unavailable_response(request, "hechos_delictivos")
    hecho = get_object_or_404(
        HechoDelictivo.objects.select_related("creado_por", "articulo").prefetch_related(
            "autor", "noticias"
        ),
        id=id,
    )
    if request.user.userprofile.rol not in [Roles.REDACCION, Roles.ADMIN]:
        return render(request, "403.html", status=403)
    if (
        request.user.userprofile.rol == Roles.REDACCION
        and hecho.creado_por not in [request.user, None]
    ):
        return render(request, "403.html", status=403)

    context = {
        "hecho": hecho,
        "autores": hecho.autor.all(),
        "noticias": hecho.noticias.all(),
        "creado_por": hecho.creado_por.username if hecho.creado_por else "No registrado",
        "fecha_registro": hecho.fecha.strftime("%d/%m/%Y"),
        "generado_en": localtime(now()).strftime("%d/%m/%Y %H:%M"),
        "articulo": hecho.articulo,
    }

    html = render_to_string("hecho_delictivo_pdf.html", context)
    pdf_file = HTML(string=html, base_url=request.build_absolute_uri()).write_pdf()

    fecha_str = hecho.fecha.strftime("%Y-%m-%d")
    categoria = slugify(hecho.get_categoria_display() or "hecho")
    nombre_archivo = f"HechoDelictivo_{categoria}_{fecha_str}.pdf"

    response = HttpResponse(pdf_file, content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="{nombre_archivo}"'
    return response


# -------------------- API ADICIONAL (No REST Framework) --------------------

# Nota: Esta función `links_list` es muy similar a `api_links` y al ViewSet. 
# Si estás usando el ViewSet para el frontend, esta función podría ser redundante.
@login_required
def links_list(request):
    queryset = LinkRelevante.objects.filter(cargado_por=request.user)

    fecha_inicio = request.GET.get("fecha_inicio")
    fecha_fin = request.GET.get("fecha_fin")
    diario_id = request.GET.get("diario_id")
    estado = request.GET.get("estado")
    estado_in = request.GET.get("estado__in")
    fuente = request.GET.get("fuente")

    if fecha_inicio:
        queryset = queryset.filter(fecha_carga__date__gte=parse_date(fecha_inicio))
    if fecha_fin:
        queryset = queryset.filter(fecha_carga__date__lte=parse_date(fecha_fin))
    if diario_id and diario_id.isdigit():
        queryset = queryset.filter(diario_digital_id=int(diario_id))
    if estado:
        queryset = queryset.filter(estado=estado)
    if estado_in:
        estados = estado_in.split(",")
        queryset = queryset.filter(estado__in=estados)
    if fuente == "red_social":
        queryset = queryset.none()
    elif fuente == "diario":
        queryset = queryset.filter(diario_digital__isnull=False)

    serializer = LinkRelevanteSerializer(
        queryset.distinct().order_by("-fecha_carga"), many=True
    )
    return JsonResponse(serializer.data, safe=False)


# -------------------- EXPORTACIÓN PDF INFORME INDIVIDUAL --------------------


@login_required
def exportar_informe_pdf(request, informe_id):
    if HTML is None:
        return _weasyprint_unavailable_response(request, "informes")
    informe = get_object_or_404(InformeIndividual, id=informe_id)

    if request.user.userprofile.rol not in [Roles.ADMIN, Roles.GERENCIA, Roles.CLIENTE, Roles.INFORMES]:
        return render(request, "403.html", status=403)

    # El prefetch lo hacemos aquí si el objeto no está cargado completamente:
    informe_completo = InformeIndividual.objects.select_related("generado_por").prefetch_related(
        "alias", "telefono", "domicilio", "vehiculos", "empleadores", "vinculos", "articulos"
    ).get(id=informe_id)
    
    html_string = render_to_string("informes/informe_pdf.html", {"informe": informe_completo})
    
    try:
        pdf = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf()
    except Exception as e:
        logger.error(f"Error al generar PDF del informe ID {informe_id}: {e}")
        messages.error(request, "Error al generar el PDF del informe.")
        return redirect('informes') # Redirigir a una página segura

    nombre_archivo = f"Informe_{informe.apellido}_{informe.nombre}_{datetime.now().strftime('%Y%m%d')}.pdf".replace(" ", "_")
    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f"attachment; filename={nombre_archivo}"
    return response
